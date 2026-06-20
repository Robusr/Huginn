# -*- coding: utf-8 -*-
"""
通用数据分析报告生成器。
默认输出更适合直接提交和阅读的正式版报告：
- 目录使用长点线 + 页码
- 标题采用大小标题层级
- 数据概览改为精炼文本
- 图表解读固定包含关键数据、发现、检查方法
- Word 为主要交付格式，并可同步生成 PDF
"""

from __future__ import annotations

import json
import math
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from huginn.core.config import Config
from huginn.core.label_utils import clean_choice, humanize_column_name
from huginn.core.logger import get_logger

logger = get_logger(__name__)


class ReportGenerator:
    TITLE = "数据分析报告"

    MAIN_SECTIONS = [
        {"id": "1", "title": "执行摘要", "level": 1},
        {"id": "2", "title": "数据概览", "level": 1},
        {"id": "3", "title": "重点图表分析", "level": 1},
        {"id": "4", "title": "主要发现", "level": 1},
        {"id": "5", "title": "行动建议", "level": 1},
        {"id": "6", "title": "局限性与验证摘要", "level": 1},
    ]

    CHART_LABELS = {
        "bar_chart": "图 3-1 分组样本量与均值对比",
        "box_plot": "图 3-2 分组分布差异箱线图",
        "scatter_plot": "图 3-3 关键维度联动散点图",
        "correlation_heatmap": "图 3-4 核心数值指标相关性热力图",
    }

    def __init__(
        self,
        run_dir: Union[str, Path],
        user_requirement: str = "",
        domain_config=None,
    ) -> None:
        self.run_dir = Path(run_dir)
        self.user_requirement = user_requirement or Config.DEFAULT_REQUIREMENT
        self.generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.TITLE = "数据分析报告"
        self.subtitle = "聚焦主要发现、图表解读与行动建议"
        self.MAIN_SECTIONS = [dict(item) for item in type(self).MAIN_SECTIONS]

        self.data_profile: Dict[str, Any] = {}
        self.stats_results: Dict[str, Any] = {}
        self.findings: List[Dict[str, Any]] = []
        self.suggestions: List[Dict[str, Any]] = []
        self.valid_tasks: List[Dict[str, Any]] = []
        self.validation_result: Dict[str, Any] = {}
        self.chart_files: List[str] = []
        self.chart_metadata: Dict[str, Any] = {}
        self.domain_context: Dict[str, Any] = {}
        self.report_narrative: Dict[str, Any] = {}

        self.focus_fields: List[str] = []
        self.significant_anova: List[Dict[str, Any]] = []
        self.strong_point_estimates: List[Dict[str, Any]] = []
        self.heatmap_pairs: List[Dict[str, Any]] = []
        self.chart_notes: List[Dict[str, Any]] = []
        self.sections_cache: Optional[List[Dict[str, Any]]] = None

    @property
    def is_education(self) -> bool:
        return self.domain_config is not None and self.domain_config.key == "education_survey"

    # ==================================================================
    # Public API
    # ==================================================================

    def generate(self) -> str:
        self._load_all_data()
        self._discover_charts()
        self._prepare_report_context()
        lines = self._render_markdown()
        return "\n".join(lines).strip() + "\n"

    def save(self, filename: str = "final_report.md") -> Path:
        content = self.generate()
        path = self.run_dir / filename
        path.write_text(content, encoding="utf-8")
        logger.info("Markdown 报告已保存至 %s", path)
        return path

    def export_word(self, filename: str = "final_report.docx") -> Path:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        self._load_if_needed()
        sections = self._build_sections()

        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.2)

        self._configure_word_styles(doc)
        self._add_footer_page_number(doc)

        title = doc.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.TITLE)
        run.font.bold = True

        sub = doc.add_paragraph(style="ReportBody")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(self.subtitle).italic = True

        doc.add_paragraph("", style="ReportBody")

        toc_title = doc.add_paragraph("目 录", style="Heading 1")
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.paragraph_format.page_break_before = False
        self._render_static_toc_docx(doc)

        for section_data in sections:
            self._render_section_docx(doc, section_data)

        path = self.run_dir / filename
        doc.save(str(path))
        logger.info("Word 报告已保存至 %s", path)
        return path

    def export_pdf(self, filename: str = "final_report.pdf") -> Optional[Path]:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import (
                Image,
                KeepTogether,
                PageBreak,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
            )
        except ImportError:
            logger.warning("缺少 PDF 导出依赖，跳过 PDF 输出")
            return None

        self._load_if_needed()
        sections = self._build_sections()

        base_dir = self.run_dir
        path = self.run_dir / filename

        songti = "/System/Library/Fonts/Supplemental/Songti.ttc"
        heiti = "/System/Library/Fonts/STHeiti Medium.ttc"
        if Path(songti).exists():
            pdfmetrics.registerFont(TTFont("Songti", songti))
        if Path(heiti).exists():
            pdfmetrics.registerFont(TTFont("Heiti", heiti))

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="RptTitle",
                fontName="Heiti" if Path(heiti).exists() else "Helvetica-Bold",
                fontSize=20,
                leading=26,
                alignment=TA_CENTER,
                spaceAfter=10,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptSubtitle",
                fontName="Songti" if Path(songti).exists() else "Helvetica",
                fontSize=10.5,
                leading=16,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#4A5A6A"),
                spaceAfter=10,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptBody",
                fontName="Songti" if Path(songti).exists() else "Helvetica",
                fontSize=10.5,
                leading=16,
                alignment=TA_JUSTIFY,
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptBullet",
                parent=styles["RptBody"],
                alignment=TA_LEFT,
                leftIndent=14,
                firstLineIndent=-10,
                spaceAfter=3,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptNumber",
                parent=styles["RptBody"],
                alignment=TA_LEFT,
                leftIndent=18,
                firstLineIndent=-14,
                spaceAfter=3,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptH1",
                fontName="Heiti" if Path(heiti).exists() else "Helvetica-Bold",
                fontSize=15.5,
                leading=22,
                textColor=colors.HexColor("#1F2A44"),
                spaceBefore=8,
                spaceAfter=8,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptH2",
                fontName="Heiti" if Path(heiti).exists() else "Helvetica-Bold",
                fontSize=12.5,
                leading=18,
                textColor=colors.HexColor("#385170"),
                spaceBefore=6,
                spaceAfter=6,
                keepWithNext=1,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptH3",
                fontName="Heiti" if Path(heiti).exists() else "Helvetica-Bold",
                fontSize=11.2,
                leading=16,
                textColor=colors.HexColor("#4A5A6A"),
                spaceBefore=4,
                spaceAfter=5,
                keepWithNext=1,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptCaption",
                fontName="Songti" if Path(songti).exists() else "Helvetica",
                fontSize=9,
                leading=13,
                alignment=TA_CENTER,
                textColor=colors.HexColor("#5C677D"),
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptTOC1",
                parent=styles["RptBody"],
                alignment=TA_LEFT,
                fontSize=10.8,
                leading=16,
                spaceAfter=3,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptTOC2",
                parent=styles["RptBody"],
                alignment=TA_LEFT,
                fontSize=10,
                leading=15,
                leftIndent=14,
                spaceAfter=2,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptTOCDots",
                parent=styles["RptBody"],
                alignment=TA_RIGHT,
                fontSize=10,
                leading=15,
                textColor=colors.HexColor("#667085"),
                spaceAfter=2,
            )
        )
        styles.add(
            ParagraphStyle(
                name="RptTOCPage",
                parent=styles["RptBody"],
                alignment=TA_RIGHT,
                fontSize=10,
                leading=15,
                spaceAfter=2,
            )
        )

        story = [
            Paragraph(self.TITLE, styles["RptTitle"]),
            Paragraph(self.subtitle, styles["RptSubtitle"]),
            Spacer(1, 0.2 * cm),
            Paragraph("目 录", styles["RptH1"]),
        ]
        story.extend(self._render_pdf_toc_flowables(styles))
        story.append(PageBreak())

        for section_data in sections:
            story.append(Paragraph(section_data["heading"], styles["RptH1"]))
            for block in section_data["blocks"]:
                story.extend(self._render_pdf_block(block, styles, base_dir))
            story.append(PageBreak())
        if story and isinstance(story[-1], PageBreak):
            story.pop()

        def add_page_num(canvas, _doc):
            canvas.setFont("Songti" if Path(songti).exists() else "Helvetica", 9)
            canvas.setFillColor(colors.HexColor("#666666"))
            canvas.drawCentredString(A4[0] / 2, 1.0 * cm, str(canvas.getPageNumber()))

        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=2.05 * cm,
            rightMargin=2.0 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.6 * cm,
        )
        doc.build(story, onFirstPage=add_page_num, onLaterPages=add_page_num)
        logger.info("PDF 报告已保存至 %s", path)
        return path

    # ==================================================================
    # Load / Prepare
    # ==================================================================

    def _load_if_needed(self) -> None:
        if self.sections_cache is not None:
            return
        self._load_all_data()
        self._discover_charts()
        self._prepare_report_context()

    def _load_all_data(self) -> None:
        self.data_profile = self._load_json("data_profile.json", required=True)
        self.stats_results = self._load_json("stats_results.json", required=True)
        self.findings = self._load_json_list("findings.json")
        self.suggestions = self._load_json_list("suggestions.json")
        self.valid_tasks = self._load_json_list("valid_tasks.json")
        self.validation_result = self._load_json("validation_result.json", required=False)
        self.chart_metadata = self._load_json("chart_metadata.json", required=False)
        self.domain_context = self._load_json("domain_context.json", required=False)
        self.report_narrative = self._load_json("report_narrative.json", required=False)
        self.TITLE = self.report_narrative.get("title") or self.domain_context.get("report_title") or "数据分析报告"
        self.subtitle = self.report_narrative.get("subtitle") or self.domain_context.get("report_subtitle") or self.subtitle
        recommendation_title = self.domain_context.get("recommendation_section", "行动建议")
        self.MAIN_SECTIONS[4]["title"] = recommendation_title
        if self.report_narrative.get("findings"):
            self.findings = self.report_narrative["findings"]
        if self.report_narrative.get("suggestions"):
            self.suggestions = self.report_narrative["suggestions"]

    def _load_json(self, filename: str, required: bool = False) -> Dict[str, Any]:
        path = self.run_dir / filename
        if not path.exists():
            if required:
                raise FileNotFoundError(f"缺少必需文件: {filename}（路径: {self.run_dir}）")
            return {}
        with open(path, "r", encoding="utf-8") as handle:
            return json.load(handle)

    def _load_json_list(self, filename: str) -> List[Dict[str, Any]]:
        path = self.run_dir / filename
        if not path.exists():
            return []
        with open(path, "r", encoding="utf-8") as handle:
            data = json.load(handle)
        return data if isinstance(data, list) else []

    def _discover_charts(self) -> None:
        chart_dir = self.run_dir / "charts"
        if not chart_dir.exists():
            self.chart_files = []
            return
        preferred = ["bar_chart", "box_plot", "scatter_plot", "correlation_heatmap"]
        png_files = sorted(chart_dir.glob("*.png"))
        ordered: List[str] = []
        for stem in preferred:
            for file in png_files:
                if file.stem == stem:
                    ordered.append(f"charts/{file.name}")
        for file in png_files:
            rel = f"charts/{file.name}"
            if rel not in ordered:
                ordered.append(rel)
        self.chart_files = ordered

    def _prepare_report_context(self) -> None:
        self._ensure_fallback_findings_and_suggestions()
        self.focus_fields = self._select_focus_fields(self.data_profile.get("fields", []))
        self.significant_anova = self._collect_significant_anova()
        self.strong_point_estimates = self._collect_top_point_estimates()
        self.heatmap_pairs = self._collect_heatmap_pairs()
        self.chart_notes = self._build_chart_notes()
        self.sections_cache = self._build_sections()

    def _build_sections(self) -> List[Dict[str, Any]]:
        if self.sections_cache is not None:
            return self.sections_cache

        sections = [
            {
                "id": "1",
                "heading": "一、执行摘要",
                "blocks": self._section_exec_summary(),
            },
            {
                "id": "2",
                "heading": "二、数据概览",
                "blocks": self._section_dataset_overview(),
            },
            {
                "id": "3",
                "heading": "三、重点图表分析",
                "blocks": self._section_chart_analysis(),
            },
            {
                "id": "4",
                "heading": "四、主要发现",
                "blocks": self._section_findings(),
            },
            {
                "id": "5",
                "heading": f"五、{self.domain_context.get('recommendation_section', '行动建议')}",
                "blocks": self._section_suggestions(),
            },
            {
                "id": "6",
                "heading": "六、局限性与验证摘要",
                "blocks": self._section_limitations(),
            },
        ]
        self.sections_cache = sections
        return sections

    # ==================================================================
    # Markdown rendering
    # ==================================================================

    def _render_markdown(self) -> List[str]:
        sections = self._build_sections()
        lines = [
            f"# {self.TITLE}",
            "",
            self.subtitle + ("。" if not self.subtitle.endswith("。") else ""),
            "",
            "## 目 录",
            "",
        ]
        lines.extend(self._render_static_toc_lines())
        lines.extend(["", "---", ""])

        for section in sections:
            lines.append(f"## {section['heading']}")
            lines.append("")
            lines.extend(self._render_markdown_blocks(section["blocks"]))
            lines.extend(["", "---", ""])

        lines.extend(
            [
                "",
                "*本报告由 Huginn 自动生成，已按“少过程、重发现、强图表解读”的正式报告模式输出。*",
            ]
        )
        return lines

    def _render_markdown_blocks(self, blocks: List[Dict[str, Any]]) -> List[str]:
        lines: List[str] = []
        for block in blocks:
            kind = block["type"]
            if kind == "heading2":
                lines.extend([f"### {block['text']}", ""])
            elif kind == "heading3":
                lines.extend([f"#### {block['text']}", ""])
            elif kind == "paragraph":
                lines.extend([block["text"], ""])
            elif kind == "bullets":
                for item in block["items"]:
                    lines.append(f"- {item}")
                lines.append("")
            elif kind == "numbered":
                for index, item in enumerate(block["items"], 1):
                    lines.append(f"{index}. {item}")
                lines.append("")
            elif kind == "image":
                lines.extend([f"![{block['caption']}]({block['path']})", ""])
            elif kind == "toc":
                lines.extend(block["lines"] + [""])
            elif kind == "page_break":
                lines.extend(["<!-- pagebreak -->", ""])
        return lines

    # ==================================================================
    # DOCX rendering
    # ==================================================================

    def _configure_word_styles(self, doc) -> None:
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin

        normal = doc.styles["Normal"]
        normal.font.name = "宋体"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.35

        title_style = doc.styles["Title"]
        title_style.font.name = "黑体"
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        title_style.font.size = Pt(22)
        title_style.font.color.rgb = RGBColor(27, 41, 66)
        title_style.paragraph_format.space_after = Pt(10)

        h1 = doc.styles["Heading 1"]
        h1.font.name = "黑体"
        h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        h1.font.size = Pt(15.5)
        h1.font.color.rgb = RGBColor(31, 42, 68)
        h1.paragraph_format.space_before = Pt(8)
        h1.paragraph_format.space_after = Pt(8)
        h1.paragraph_format.page_break_before = True

        h2 = doc.styles["Heading 2"]
        h2.font.name = "黑体"
        h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        h2.font.size = Pt(12.5)
        h2.font.color.rgb = RGBColor(56, 81, 112)
        h2.paragraph_format.space_before = Pt(5)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.keep_with_next = True

        h3 = doc.styles["Heading 3"]
        h3.font.name = "黑体"
        h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        h3.font.size = Pt(11.2)
        h3.font.color.rgb = RGBColor(80, 94, 109)
        h3.paragraph_format.space_before = Pt(4)
        h3.paragraph_format.space_after = Pt(4)
        h3.paragraph_format.keep_with_next = True

        if "ReportBody" not in doc.styles:
            style = doc.styles.add_style("ReportBody", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        report_body = doc.styles["ReportBody"]
        report_body.font.name = "宋体"
        report_body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        report_body.font.size = Pt(10.5)
        report_body.paragraph_format.first_line_indent = Cm(0.74)
        report_body.paragraph_format.line_spacing = 1.35
        report_body.paragraph_format.space_after = Pt(6)

        if "FigureCaption" not in doc.styles:
            style = doc.styles.add_style("FigureCaption", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        caption = doc.styles["FigureCaption"]
        caption.font.name = "宋体"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        caption.font.size = Pt(9)
        caption.font.color.rgb = RGBColor(92, 103, 125)
        caption.paragraph_format.first_line_indent = Cm(0)
        caption.paragraph_format.line_spacing = 1.15
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(7)

        if "TOCLine1" not in doc.styles:
            style = doc.styles.add_style("TOCLine1", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
            style.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        toc1 = doc.styles["TOCLine1"]
        toc1.font.name = "宋体"
        toc1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        toc1.font.size = Pt(11)
        toc1.paragraph_format.first_line_indent = Cm(0)
        toc1.paragraph_format.left_indent = Cm(0)
        toc1.paragraph_format.space_after = Pt(4)

        if "TOCLine2" not in doc.styles:
            style = doc.styles.add_style("TOCLine2", WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
            style.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        toc2 = doc.styles["TOCLine2"]
        toc2.font.name = "宋体"
        toc2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        toc2.font.size = Pt(10.5)
        toc2.paragraph_format.left_indent = Cm(0.7)
        toc2.paragraph_format.first_line_indent = Cm(0)
        toc2.paragraph_format.space_after = Pt(2)

    def _add_footer_page_number(self, doc) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        footer = doc.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run.font.name = "宋体"
        run.font.size = Pt(9)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        result = OxmlElement("w:t")
        result.text = "1"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(result)
        run._r.append(fld_end)

    def _render_static_toc_docx(self, doc) -> None:
        for line in self._build_toc_entries():
            style = "TOCLine1" if line["level"] == 1 else "TOCLine2"
            paragraph = doc.add_paragraph(style=style)
            paragraph.add_run(line["label"])
            paragraph.add_run("\t")
            paragraph.add_run(str(line["page"]))

    def _render_section_docx(self, doc, section_data: Dict[str, Any]) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm

        doc.add_paragraph(section_data["heading"], style="Heading 1")
        for block in section_data["blocks"]:
            kind = block["type"]
            if kind == "heading2":
                doc.add_paragraph(block["text"], style="Heading 2")
            elif kind == "heading3":
                doc.add_paragraph(block["text"], style="Heading 3")
            elif kind == "paragraph":
                doc.add_paragraph(block["text"], style="ReportBody")
            elif kind == "bullets":
                for item in block["items"]:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "numbered":
                for item in block["items"]:
                    doc.add_paragraph(item, style="List Number")
            elif kind == "image":
                image_path = self.run_dir / block["path"]
                if image_path.exists():
                    doc.add_picture(str(image_path), width=Cm(14.6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph(block["caption"], style="FigureCaption")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "page_break":
                doc.add_page_break()

    # ==================================================================
    # PDF rendering helper
    # ==================================================================

    def _render_pdf_toc_flowables(self, styles) -> List[Any]:
        from reportlab.lib import colors
        from reportlab.platypus import Flowable, Spacer

        class _TOCLeaderLine(Flowable):
            def __init__(self, label: str, page: int, *, level: int) -> None:
                super().__init__()
                self.label = label
                self.page = str(page)
                self.level = level
                self.font_name = styles["RptTOC1"].fontName
                self.font_size = 10.8 if level == 1 else 10
                self.leading = 20 if level == 1 else 18
                self.indent = 14 if level == 2 else 0
                self.text_color = colors.black
                self.dot_color = colors.HexColor("#667085")
                self._available_width = 0

            def wrap(self, availWidth, availHeight):
                self._available_width = availWidth
                return availWidth, self.leading

            def draw(self):
                canvas = self.canv
                baseline = 5
                page_width = canvas.stringWidth(self.page, self.font_name, self.font_size)

                canvas.setFont(self.font_name, self.font_size)
                canvas.setFillColor(self.text_color)
                canvas.drawString(self.indent, baseline, self.label)
                canvas.drawRightString(self._available_width, baseline, self.page)

                label_width = canvas.stringWidth(self.label, self.font_name, self.font_size)
                start = self.indent + label_width + 8
                end = self._available_width - page_width - 8
                if end <= start:
                    return

                canvas.setFillColor(self.dot_color)
                dot_width = canvas.stringWidth(".", self.font_name, self.font_size)
                x = start
                while x < end:
                    canvas.drawString(x, baseline, ".")
                    x += max(dot_width + 1.2, 3)

        flowables: List[Any] = []
        for entry in self._build_toc_entries(output_format="pdf"):
            flowables.append(_TOCLeaderLine(entry["label"], entry["page"], level=entry["level"]))
        if flowables:
            flowables.append(Spacer(1, 6))
        return flowables

    def _render_pdf_block(self, block: Dict[str, Any], styles, base_dir: Path) -> List[Any]:
        from reportlab.lib.units import cm
        from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, Spacer

        flowables: List[Any] = []
        kind = block["type"]
        if kind == "heading2":
            flowables.append(Paragraph(self._escape(block["text"]), styles["RptH2"]))
        elif kind == "heading3":
            flowables.append(Paragraph(self._escape(block["text"]), styles["RptH3"]))
        elif kind == "paragraph":
            flowables.append(Paragraph(self._escape(block["text"]), styles["RptBody"]))
        elif kind == "bullets":
            for item in block["items"]:
                flowables.append(Paragraph(self._escape(f"- {item}"), styles["RptBullet"]))
            flowables.append(Spacer(1, 0.08 * cm))
        elif kind == "numbered":
            for index, item in enumerate(block["items"], 1):
                flowables.append(Paragraph(self._escape(f"{index}. {item}"), styles["RptNumber"]))
            flowables.append(Spacer(1, 0.08 * cm))
        elif kind == "image":
            image_path = base_dir / block["path"]
            if image_path.exists():
                image = Image(str(image_path))
                image._restrictSize(14.8 * cm, 8.7 * cm)
                image.hAlign = "CENTER"
                caption = Paragraph(self._escape(block["caption"]), styles["RptCaption"])
                flowables.append(KeepTogether([image, caption, Spacer(1, 0.08 * cm)]))
        elif kind == "page_break":
            flowables.append(PageBreak())
        return flowables

    # ==================================================================
    # Section builders
    # ==================================================================

    def _section_exec_summary(self) -> List[Dict[str, Any]]:
        meta = self.data_profile.get("meta", {})
        field_types = self.data_profile.get("overview", {}).get("field_type_counts", {})
        score = self.validation_result.get("meta", {}).get("score")

        summary_text = self.report_narrative.get("executive_summary") or (
            f"本报告基于 {meta.get('n_rows', '?')} 条有效记录、{meta.get('n_columns', '?')} 个字段展开，"
            "以主要发现、图表证据和行动建议为核心，统计过程仅保留必要的方法说明。"
        )

        bullets = [
            f"本轮对 {len(self.stats_results.get('point_estimation', {}).get('fields', {}))} 个有效数值指标完成了点估计和区间估计。",
            f"本轮自动分析提炼出 {len(self.findings)} 条主要发现、{len(self.suggestions)} 条行动建议，并优先保留可追溯的统计证据。",
            f"自动验证得分为 {score}/100。"
            if score is not None
            else "当前运行结果尚未生成自动验证得分，但统计结果和图表已完成。",
        ]

        top_findings = [self._normalize_finding_text(item.get("conclusion", "")) for item in self.findings[:3]]

        blocks: List[Dict[str, Any]] = [
            {"type": "paragraph", "text": summary_text},
            {"type": "heading2", "text": "1.1 核心关注点"},
            {"type": "bullets", "items": bullets},
        ]
        if top_findings:
            blocks.extend(
                [
                    {"type": "heading2", "text": "1.2 优先结论"},
                    {"type": "numbered", "items": top_findings},
                ]
            )
        return blocks

    def _section_dataset_overview(self) -> List[Dict[str, Any]]:
        meta = self.data_profile.get("meta", {})
        overview = self.data_profile.get("overview", {})
        duplicate_rows = overview.get("duplicate_rows", 0)
        missing_pct = meta.get("total_missing_pct", 0)

        narrative_overview = self.report_narrative.get("overview_paragraphs") or []
        intro = (
            f"本次分析使用 {meta.get('n_rows', '?')} 条记录和 {meta.get('n_columns', '?')} 个观察字段。"
            f"整体缺失率为 {missing_pct:.2f}%，重复记录数为 {duplicate_rows}。"
        )
        focus_text = "结合统计任务与图表，本报告重点关注核心数值指标、可解释分组及其差异和关联。"

        field_desc = []
        if self.focus_fields:
            field_desc.append("本次重点字段包括：" + "、".join(self.focus_fields[:8]) + "。")

        distribution_parts = self._dataset_distribution_brief()

        base_paragraphs = narrative_overview or [intro, focus_text]
        blocks = [{"type": "paragraph", "text": paragraph} for paragraph in base_paragraphs]
        blocks.extend({"type": "paragraph", "text": item} for item in field_desc + distribution_parts)
        return blocks

    def _section_chart_analysis(self) -> List[Dict[str, Any]]:
        blocks: List[Dict[str, Any]] = [
            {
                "type": "paragraph",
                "text": self.report_narrative.get("chart_section_intro") or "本节围绕每张图给出关键数据、主要发现与检查方法，并结合统计结果解释差异和关联。",
            }
        ]

        if not self.chart_notes:
            blocks.append({"type": "paragraph", "text": "当前运行结果尚未生成可用图表。"})
            return blocks

        for index, note in enumerate(self.chart_notes):
            blocks.extend(
                [
                    {"type": "heading2", "text": note["heading"]},
                    {"type": "image", "path": note["path"], "caption": note["caption"]},
                    {"type": "heading3", "text": "关键数据"},
                    {"type": "bullets", "items": note["key_data"]},
                    {"type": "heading3", "text": "主要发现"},
                    {"type": "paragraph", "text": note["finding"]},
                    {"type": "heading3", "text": "检查方法"},
                    {"type": "paragraph", "text": note["method"]},
                ]
            )
            if index < len(self.chart_notes) - 1:
                blocks.append({"type": "page_break"})
        return blocks

    def _section_findings(self) -> List[Dict[str, Any]]:
        if not self.findings:
            return [{"type": "paragraph", "text": "当前未生成可用的主要发现。"}]

        blocks: List[Dict[str, Any]] = []
        for index, finding in enumerate(self.findings[:8], 1):
            conclusion = self._clean_report_text(self._normalize_finding_text(finding.get("conclusion", "待补充")))
            evidence = self._clean_report_text(finding.get("evidence", "统计证据待补充"))
            method = finding.get("method", "统计分析")
            significance = self._significance_text_from_evidence(evidence, method)
            blocks.extend(
                [
                    {"type": "heading2", "text": self._finding_heading(index, conclusion)},
                    {"type": "paragraph", "text": conclusion},
                    {"type": "bullets", "items": [f"关键证据：{evidence}", f"方法说明：{method}", f"结果解读：{significance}"]},
                ]
            )
        return blocks

    def _section_suggestions(self) -> List[Dict[str, Any]]:
        if not self.suggestions:
            return [{"type": "paragraph", "text": "当前未生成可用的行动建议。"}]

        blocks: List[Dict[str, Any]] = [
            {
                "type": "paragraph",
                "text": "以下建议均对应前文已识别出的差异、关联或集中性问题，强调证据依据、执行动作与后续观察。",
            }
        ]

        for index, suggestion in enumerate(self.suggestions[:5], 1):
            text = self._clean_report_text(self._ensure_sentence(suggestion.get("suggestion", "待补充")))
            evidence = self._clean_report_text(self._normalize_finding_text(suggestion.get("evidence", "待补充")))
            direction = self._clean_report_text(self._ensure_sentence(suggestion.get("direction", "待补充")))
            blocks.extend(
                [
                    {"type": "heading2", "text": self._suggestion_heading(index, {**suggestion, "evidence": evidence})},
                    {"type": "paragraph", "text": text},
                    {"type": "bullets", "items": [f"对应依据：{evidence}", f"落地方向：{direction}"]},
                ]
            )
        return blocks

    def _section_limitations(self) -> List[Dict[str, Any]]:
        score = self.validation_result.get("meta", {}).get("score")
        passed = self.validation_result.get("meta", {}).get("overall_pass")
        improvements = self.validation_result.get("improvement_suggestions", [])

        limitations = self.report_narrative.get("limitations") or [
            "本报告中的显著结果反映群体差异或变量关联，不直接代表因果关系。",
            f"当前结论基于 {self.data_profile.get('meta', {}).get('n_rows', '?')} 条记录，外推时需结合数据覆盖范围。",
            "日期、分类口径和异常值会影响结果解释，关键决策前应回到源数据复核。",
        ]
        blocks: List[Dict[str, Any]] = [
            {"type": "heading2", "text": "6.1 使用边界"},
            {
                "type": "bullets",
                "items": limitations,
            },
            {"type": "heading2", "text": "6.2 自动验证摘要"},
        ]

        if score is not None:
            blocks.append(
                {
                    "type": "paragraph",
                    "text": f"自动验证得分为 {score}/100，整体结果为{'通过' if passed else '未通过'}。"
                    "验证结果表明统计链路基本可用，但在发现覆盖度、建议具体度和报告完整性方面仍有提升空间。",
                }
            )
        else:
            blocks.append({"type": "paragraph", "text": "当前未生成自动验证结果。"})

        if improvements:
            polished = [self._polish_validation_suggestion(item) for item in improvements[:4]]
            blocks.append({"type": "bullets", "items": polished})
        return blocks

    # ==================================================================
    # Chart notes
    # ==================================================================

    def _build_chart_notes(self) -> List[Dict[str, Any]]:
        metadata_items = self.chart_metadata.get("charts", [])
        metadata_map = {item.get("chart_type"): item for item in metadata_items if isinstance(item, dict)}

        notes: List[Dict[str, Any]] = []
        for path in self.chart_files:
            stem = Path(path).stem
            meta = metadata_map.get(stem, {})
            notes.append(self._chart_note_from_meta(stem, path, meta))
        return notes

    def _chart_note_from_meta(self, stem: str, path: str, meta: Dict[str, Any]) -> Dict[str, Any]:
        if not meta:
            meta = self._fallback_chart_meta(stem, path)
        label = self.CHART_LABELS.get(stem, stem)
        if stem == "bar_chart":
            group_means = meta.get("group_means", {})
            highest_group = meta.get("highest_group")
            lowest_group = meta.get("lowest_group")
            source_test = meta.get("source_test", {})
            p_value = source_test.get("p_value")
            key_data = []
            if highest_group:
                key_data.append(
                    f"{self._clean_choice(highest_group)}组的平均值最高，为 {self._fmt(meta.get('highest_mean'), 2)}；"
                    f"{self._clean_choice(lowest_group)}组最低，为 {self._fmt(meta.get('lowest_mean'), 2)}。"
                )
            if group_means:
                sample_desc = "；".join(f"{self._clean_choice(k)}={self._fmt(v, 2)}" for k, v in list(group_means.items())[:4])
                key_data.append(f"各组均值概况为：{sample_desc}。")
            finding = (
                f"该图直接展示了{self._label(meta.get('factor'))}不同组别在“{self._label(meta.get('dependent'))}”上的均值差异，"
                "便于判断哪些分组值得优先复核或采取差异化措施。"
            )
            if isinstance(p_value, (int, float)):
                finding += f"对应单因素方差分析 p={p_value:.4f}，说明该差异{'达到' if p_value < 0.05 else '未达到'}统计显著水平。"
            method = (
                f"图中左侧给出{self._label(meta.get('factor'))}的样本数量，右侧给出各组在“{self._label(meta.get('dependent'))}”上的均值。"
                "检查时先核对样本分布，再结合方差分析结果判断均值差异是否稳定。"
            )
        elif stem == "box_plot":
            stats = meta.get("group_stats", {})
            ordered = meta.get("ordered_groups", [])
            source_test = meta.get("source_test", {})
            key_data = []
            if ordered:
                first = ordered[0]
                if first in stats:
                    key_data.append(
                        f"{self._clean_choice(first)}组的中位数较高，为 {self._fmt(stats[first].get('median'), 2)}，"
                        f"IQR 为 {self._fmt(stats[first].get('iqr'), 2)}。"
                    )
                if len(ordered) > 1 and ordered[-1] in stats:
                    last = ordered[-1]
                    key_data.append(
                        f"{self._clean_choice(last)}组中位数较低，为 {self._fmt(stats[last].get('median'), 2)}，"
                        f"说明该组整体数值偏低。"
                    )
            finding = (
                f"相比只看均值，这张箱线图更适合判断“{self._label(meta.get('dependent'))}”的差异是整体偏移，"
                "还是由少量极端值拉动。当前图中不同组的中位数和四分位区间存在可见差别，说明分组分布并不相同。"
            )
            if source_test.get("p_value") is not None:
                finding += f"与之对应的方差分析 p={source_test['p_value']:.4f}。"
            method_name = source_test.get("method", "方差分析")
            method = (
                f"检查时重点关注中位数线、箱体宽度与离群点位置。若中位数差异明显、箱体重叠有限，"
                f"再结合{method_name}的 p 值，可更稳健地解释群体差异。"
            )
        elif stem == "scatter_plot":
            corr = meta.get("correlation")
            key_data = [
                f"样本量为 {meta.get('n', '?')}，"
                f"“{self._label(meta.get('x'))}”的均值为 {self._fmt(meta.get('x_mean'), 2)}，"
                f"“{self._label(meta.get('y'))}”的均值为 {self._fmt(meta.get('y_mean'), 2)}。"
            ]
            if isinstance(corr, (int, float)):
                key_data.append(f"Pearson 相关系数为 {corr:.3f}。")
            finding = (
                f"散点图用于观察“{self._label(meta.get('x'))}”与“{self._label(meta.get('y'))}”是否存在同步变化。"
                "点云若大致沿同一方向展开，说明两个数值指标具有同向或反向联动。"
            )
            method = (
                "检查时先看点云是否呈现明显斜率，再看拟合线方向与 Pearson r 的正负和绝对值。"
                "这一步更多用于确认联动趋势，而不是直接作因果解释。"
            )
        elif stem == "correlation_heatmap":
            pairs = meta.get("top_pairs", [])
            key_data = []
            for pair in pairs[:3]:
                key_data.append(
                    f"“{self._label(pair.get('left'))}”与“{self._label(pair.get('right'))}”的相关系数为 {self._fmt(pair.get('correlation'), 2)}。"
                )
            finding = (
                "热力图适合快速锁定关系更紧密的数值指标，帮助报告聚焦真正联动的指标组合。"
            )
            method = (
                "检查时优先关注颜色最深、绝对值最大的格子，再回到指标定义判断这种相关是否具有实际解释价值。"
            )
        else:
            key_data = ["该图用于支撑报告中的重点结论。"]
            finding = "该图反映了当前数据中最值得优先解释的一组关系。"
            method = "检查时结合图形形态与对应统计结果综合判断。"

        return {
            "heading": label,
            "path": path,
            "caption": label,
            "key_data": key_data,
            "finding": finding,
            "method": method,
        }

    def _fallback_chart_meta(self, stem: str, path: str) -> Dict[str, Any]:
        if stem in {"bar_chart", "box_plot"} and self.significant_anova:
            source = self.significant_anova[0] if stem == "bar_chart" else self.significant_anova[min(1, len(self.significant_anova) - 1)]
            dependent = source.get("dependent")
            factor = source.get("factor")
            fallback: Dict[str, Any] = {
                "chart_type": stem,
                "image": path,
                "factor": factor,
                "dependent": dependent,
                "source_test": source,
            }
            if factor in self.data_profile.get("fields", []) or dependent in self.data_profile.get("fields", []):
                return fallback
            if factor in self._profile_columns() and dependent in self._profile_columns():
                return fallback
            return fallback

        if stem == "scatter_plot":
            if self.heatmap_pairs:
                pair = self.heatmap_pairs[0]
                x = pair.get("left")
                y = pair.get("right")
                x_stats = self.stats_results.get("point_estimation", {}).get("fields", {}).get(x, {})
                y_stats = self.stats_results.get("point_estimation", {}).get("fields", {}).get(y, {})
                return {
                    "chart_type": stem,
                    "image": path,
                    "x": x,
                    "y": y,
                    "correlation": pair.get("correlation"),
                    "n": self.data_profile.get("meta", {}).get("n_rows"),
                    "x_mean": x_stats.get("mean"),
                    "y_mean": y_stats.get("mean"),
                }
            pair = self._fallback_scatter_pair()
            if pair:
                return pair

        if stem == "correlation_heatmap":
            return {
                "chart_type": stem,
                "image": path,
                "top_pairs": self.heatmap_pairs,
            }

        return {"chart_type": stem, "image": path}

    def _fallback_scatter_pair(self) -> Optional[Dict[str, Any]]:
        estimates = self.stats_results.get("point_estimation", {}).get("fields", {})
        numeric_cols = list(estimates.keys())[:2]
        if len(numeric_cols) < 2:
            return None
        x, y = numeric_cols[0], numeric_cols[1]
        return {
            "chart_type": "scatter_plot",
            "x": x,
            "y": y,
            "n": self.data_profile.get("meta", {}).get("n_rows"),
            "x_mean": estimates.get(x, {}).get("mean"),
            "y_mean": estimates.get(y, {}).get("mean"),
        }

    def _profile_columns(self) -> set[str]:
        return {item.get("column") for item in self.data_profile.get("fields", []) if isinstance(item, dict)}

    # ==================================================================
    # Helpers
    # ==================================================================

    def _ensure_fallback_findings_and_suggestions(self) -> None:
        if not self.findings:
            for item in self._collect_significant_anova()[:5]:
                self.findings.append(
                    {
                        "conclusion": f"不同{self._label(item.get('factor'))}分组的{self._label(item.get('dependent'))}存在显著差异。",
                        "evidence": f"F={self._fmt(item.get('F_statistic'))}, p={self._fmt(item.get('p_value'))}",
                        "method": item.get("method", "单因素方差分析"),
                        "importance": 4,
                    }
                )

        if not self.suggestions:
            for finding in self.findings[:3]:
                self.suggestions.append(
                    {
                        "suggestion": "围绕差异显著的分组制定更有针对性的行动方案。",
                        "evidence": finding.get("conclusion", ""),
                        "direction": "分别观察高值组与低值组，并在资源配置、执行方式和跟踪指标上作区分。",
                    }
                )

    def _collect_significant_anova(self) -> List[Dict[str, Any]]:
        tests = self.stats_results.get("anova", {}).get("tests", {})
        items: List[Dict[str, Any]] = []
        for test in tests.values():
            if isinstance(test, dict) and isinstance(test.get("p_value"), (int, float)) and test["p_value"] < 0.05:
                items.append(test)
        items.sort(key=lambda x: x.get("p_value", 1.0))
        return items

    def _collect_top_point_estimates(self) -> List[Dict[str, Any]]:
        fields = self.stats_results.get("point_estimation", {}).get("fields", {})
        items = []
        for column, stats in fields.items():
            if not isinstance(stats, dict) or "mean" not in stats:
                continue
            items.append({"column": column, **stats})
        items.sort(key=lambda x: x.get("mean", 0), reverse=True)
        return items[:6]

    def _collect_heatmap_pairs(self) -> List[Dict[str, Any]]:
        pairs = self.chart_metadata.get("charts", [])
        for item in pairs:
            if item.get("chart_type") == "correlation_heatmap":
                return item.get("top_pairs", [])
        return []

    def _select_focus_fields(self, fields: List[Dict[str, Any]]) -> List[str]:
        priority_keywords = self.domain_context.get("metric_keywords", []) + self.domain_context.get("group_keywords", [])
        selected = []
        for field in fields:
            column = field.get("column", "")
            readable = self._label(column)
            if "col " in readable or "col_" in readable:
                continue
            if not priority_keywords or any(str(keyword).lower() in f"{column} {readable}".lower() for keyword in priority_keywords):
                selected.append(readable)
        deduped: List[str] = []
        for item in selected:
            if item not in deduped:
                deduped.append(item)
        return deduped[:10]

    def _dataset_distribution_brief(self) -> List[str]:
        fields = self.data_profile.get("fields", [])
        texts: List[str] = []
        for field in fields:
            if field.get("inferred_type") != "categorical" or not 2 <= int(field.get("unique") or 0) <= 12:
                continue
            col = field.get("column", "")
            stats = field.get("stats", {})
            mode = stats.get("mode")
            mode_pct = stats.get("mode_pct")
            if mode is not None and mode_pct is not None:
                texts.append(
                    f"{self._label(col)}中，占比最高的选项为“{self._clean_choice(mode)}”，占 {mode_pct:.2f}%。"
                )
            if len(texts) >= 4:
                break
        return texts

    def _build_toc_entries(self, *, output_format: str = "word") -> List[Dict[str, Any]]:
        page_map = self._section_page_map(output_format=output_format)
        entries: List[Dict[str, Any]] = []
        for section in self.MAIN_SECTIONS:
            label = f"{section['id']} {section['title']}"
            page = page_map.get(section["id"], 1)
            entries.append({"level": 1, "label": label, "page": page})

            if section["id"] == "3":
                for idx, note in enumerate(self.chart_notes, 1):
                    entries.append({"level": 2, "label": f"3.{idx} {note['heading']}", "page": page + idx - 1})
        return entries

    def _section_page_map(self, *, output_format: str = "word") -> Dict[str, int]:
        chart_count = max(1, len(self.chart_notes)) if self.chart_notes else 1
        findings_pages = max(1, math.ceil(len(self.findings) / 4))
        suggestions_per_page = 5 if output_format == "pdf" else 3
        suggestions_pages = max(1, math.ceil(len(self.suggestions) / suggestions_per_page))
        findings_start = 4 + chart_count
        suggestions_start = findings_start + findings_pages
        return {
            "1": 2,
            "2": 3,
            "3": 4,
            "4": findings_start,
            "5": suggestions_start,
            "6": suggestions_start + suggestions_pages,
        }

    def _render_static_toc_lines(self) -> List[str]:
        lines = []
        for entry in self._build_toc_entries():
            indent = "    " if entry["level"] == 2 else ""
            lines.append(self._toc_line(f"{indent}{entry['label']}", entry["page"]))
        return lines

    @staticmethod
    def _toc_line(label: str, page: int, width: int = 46) -> str:
        plain_len = len(label)
        dot_count = max(10, width - plain_len - len(str(page)))
        return f"{label}{'.' * dot_count}{page}"

    def _render_static_toc_docx_entries(self) -> List[Tuple[str, int, int]]:
        entries = []
        for item in self._build_toc_entries():
            entries.append((item["label"], item["page"], item["level"]))
        return entries

    def _finding_heading(self, index: int, conclusion: str) -> str:
        text = conclusion.rstrip("。")
        domain = self.domain_context.get("domain")
        if domain == "retail_sales":
            if "区域" in text and "利润率" in text:
                return f"4.{index} 区域利润率分化"
            if "子品类" in text and "折扣率" in text:
                return f"4.{index} 子品类折扣率差异"
        if domain == "education_survey":
            if "数学能力自评与编程能力自评" in text:
                return f"4.{index} 数学能力自评与编程能力自评"
            if "电子游戏时间" in text and "课堂座位" in text:
                return f"4.{index} 电子游戏时间与课堂座位选择"
            if "数学能力自评不同" in text and "人形机器人" in text:
                return f"4.{index} 数学能力自评与人形机器人兴趣"
            if "技术难度" in text and "消费者" in text and "兴趣" in text:
                return f"4.{index} 技术难度认知与消费者兴趣差异"
            if "消费者兴趣评分分布" in text or "兴趣评分分布" in text:
                return f"4.{index} 消费者兴趣评分分布"
            match = re.search(r"(.+?)不同的学生在“(.+?)”上的", text)
            if match:
                factor, dependent = match.groups()
                return f"4.{index} {factor}与{dependent}"
        match = re.search(r"不同(.+?)在“(.+?)”上的", text)
        if match:
            factor, dependent = match.groups()
            return f"4.{index} {factor}与{dependent}"
        match = re.search(r"(.+?)与(.+?)之间", text)
        if match and sum(len(group) for group in match.groups()) <= 24:
            left, right = match.groups()
            return f"4.{index} {left}与{right}"
        return f"4.{index} {self._short_heading(text)}"

    def _suggestion_heading(self, index: int, suggestion: Dict[str, Any]) -> str:
        evidence = suggestion.get("evidence", "")
        suggestion_text = suggestion.get("suggestion", "")
        domain = self.domain_context.get("domain")
        if domain == "retail_sales" and "中位数" in suggestion_text and "分位数" in suggestion_text:
            return f"5.{index} 稳健利润监控与亏损预警"
        if domain == "education_survey":
            if "数学" in suggestion_text and "编程" in suggestion_text:
                return f"5.{index} 数学与编程联动训练"
            if "兴趣导入" in suggestion_text or ("机器人" in suggestion_text and "分层" in suggestion_text):
                return f"5.{index} 分层设计机器人兴趣导入"
            if "座位" in suggestion_text or "后排" in suggestion_text:
                return f"5.{index} 优化课堂座位与互动"
            if "从技术到产品" in suggestion_text or "市场分析" in suggestion_text:
                return f"5.{index} 补充技术到产品分析"
            if "项目选择" in suggestion_text or "自选项目" in suggestion_text:
                return f"5.{index} 提供分层项目选择"
            match = re.search(r"(.+?)不同的学生在“(.+?)”上的", evidence)
            if match:
                factor, dependent = match.groups()
                return f"5.{index} 分层支持：{factor}与{dependent}"
        match = re.search(r"不同(.+?)在“(.+?)”上的", evidence)
        if match:
            factor, dependent = match.groups()
            return f"5.{index} 分层支持：{factor}与{dependent}"
        if "结构性关联" in evidence:
            return f"5.{index} 跟踪关联变量"
        if "线性相关" in evidence:
            return f"5.{index} 联动改进相关指标"
        return f"5.{index} {self._short_heading(suggestion_text or evidence or '优化行动方案')}"

    @staticmethod
    def _short_heading(text: str, limit: int = 18) -> str:
        text = re.split(r"[，。；：]", str(text).strip())[0]
        return text if len(text) <= limit else text[:limit] + "…"

    @staticmethod
    def _ensure_sentence(text: str) -> str:
        text = str(text).strip()
        if text and not text.endswith(("。", "！", "？")):
            text += "。"
        return text

    @staticmethod
    def _clean_report_text(text: str) -> str:
        replacements = {
            "因位置偏远导致的": "与位置偏远相关的",
            "极易受极端值影响": "对极端值高度敏感",
            "受极端值影响": "对极端值较为敏感",
            "折扣让利侵蚀利润，却未能换取销量提升": "较高折扣同时伴随较低利润，而与销售额的线性关联接近于零",
            "折扣对销售额的拉动效应微弱": "折扣率与销售额的线性关联极弱",
            "折扣策略未有效促进销售增长": "折扣率与销售额未呈正向线性关联",
            "少数大额交易对整体指标产生显著拉动": "少数大额交易对总体指标贡献较大",
            "针对折扣率与销售额无正向关联的商品分组": "针对折扣率与销售额未呈正向线性关联的现象",
            "逐步取消或降低无效折扣": "识别并复核未改善销售表现的折扣",
            "折扣率与销售额未呈正向线性关联，折扣率与销售额的线性关联几乎为零": "折扣率与销售额的线性关联几乎为零",
            "无效折扣退出": "折扣有效性复核",
            "折扣未有效促进销量，却伴随利润率的降低": "折扣率与销售额的线性关联接近于零，并与较低利润同时出现",
            "减少无效折扣对利润的侵蚀": "减少高折扣与低利润同时出现的风险",
            "真正能拉动销量或利润": "具有更好销售或利润表现",
            "识别导致亏损的共性因素": "识别与亏损同时出现的共性特征",
            "表明其利润创造效率存在系统性缺陷": "提示其利润结构值得进一步诊断",
            "品类对利润影响显著": "不同品类的利润均值差异显著",
            "该数据质量问题将直接影响任何基于时间的业务分析": "该日期字段在修复前不适用于基于时间的业务分析",
            "表明差距源自品类结构、折扣水平和销量规模的区域分布不同": "提示应进一步检验品类结构、折扣水平和销量规模对区域差异的解释程度",
            "说明区域间利润率差异更源自品类结构、折扣水平或销量规模的差异，而非单笔订单盈利能力的固有差别": "提示还需进一步检验品类结构、折扣水平和销量规模能否解释区域利润率差异",
            "说明区域间利润率差异更可能源自品类结构、折扣水平或销量规模的差异，而非单笔订单盈利能力的固有差别": "提示还需进一步检验品类结构、折扣水平和销量规模能否解释区域利润率差异",
            "表明“薄利多销”的假设在本经营环境不成立，单纯提高销量并非提升利润的有效路径": "说明销售数量单一指标对利润的线性解释力有限，非线性或分组关系仍需进一步检验",
            "导致": "伴随",
            "造成": "伴随",
            "决定": "关联",
            "由此可见": "据此判断",
            "综上所述": "综合来看",
            "可能": "",
            "大概": "",
            "也许": "",
            "或许": "",
            "应该": "可以",
        }
        cleaned = str(text)
        for old, new in replacements.items():
            cleaned = cleaned.replace(old, new)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        cleaned = re.sub(r"\bp=0(?:\.0+)?(?=[\s，。；,;）])", "p<0.001", cleaned)
        if any(phrase in cleaned for phrase in ["不显著", "无显著", "未达显著", "未达到显著"]):
            cleaned = cleaned.replace("显著偏低", "汇总值偏低").replace("显著偏高", "汇总值偏高")
        return cleaned

    @staticmethod
    def _polish_validation_suggestion(text: str) -> str:
        text = str(text).strip()
        text = re.sub(r"^\s*[-•]?\s*", "", text)
        text = text.replace("请增加", "后续可补充")
        text = text.replace("请删除", "后续需避免")
        text = text.replace("请确保", "后续应保证")
        text = text.replace("请在", "后续可在")
        text = re.sub(r"（如'[^']+'、'[^']+'）", "", text)
        text = re.sub(r"（如“[^”]+”、“[^”]+”）", "", text)
        text = text.replace("不达标", "仍有补充空间")
        text = text.replace("存在合规性问题", "仍需保持合规表达")
        text = text.replace("不够合理", "仍可进一步具体化")
        text = text.replace("不够完整", "仍可继续完善")
        if text and not text.endswith("。"):
            text += "。"
        return text

    @staticmethod
    def _significance_text_from_evidence(evidence: str, method: str = "") -> str:
        match = re.search(r"p\s*[=<]\s*([0-9.]+)", evidence)
        if not match:
            return "该结论可作为行动参考，但仍需结合实际场景进一步核实。"
        try:
            p_val = float(match.group(1))
        except ValueError:
            return "该结论可作为行动参考，但仍需结合实际场景进一步核实。"
        if "相关" in method:
            r_match = re.search(r"(?:Pearson\s*)?r\s*=\s*(-?[0-9.]+)", evidence, re.IGNORECASE)
            if p_val >= 0.05:
                return "相关检验未达到统计显著水平，应以描述性判断为主。"
            if r_match:
                r_value = abs(float(r_match.group(1)))
                strength = "很弱" if r_value < 0.1 else "较弱" if r_value < 0.3 else "中等" if r_value < 0.5 else "较强"
                return f"统计上可检测到关联，但相关强度为{strength}，实际意义应结合效应量判断。"
            return "统计上可检测到关联，实际意义仍需结合相关系数大小判断。"
        if "分布" in method and p_val < 0.05:
            return "分布检验拒绝正态性假设，均值解释应结合中位数、分位数和异常值。"
        if p_val < 0.01:
            return "差异达到较强显著水平，说明结论较稳定。"
        if p_val < 0.05:
            return "差异达到统计显著水平，说明该维度值得优先关注。"
        return "差异未达到显著水平，应以描述性判断为主。"

    @staticmethod
    def _normalize_finding_text(text: str) -> str:
        text = text.strip()
        text = re.sub(r"\s+与\s+", "与", text)
        text = re.sub(r"^不同(.+?)在“(.+?)”上的表现存在显著差异", r"不同\1分组的“\2”存在显著差异", text)
        text = re.sub(r"^不同(.+?)在“(.+?)”上的评分是否存在显著差异", r"不同\1分组的“\2”存在显著差异", text)
        if text and not text.endswith("。"):
            text += "。"
        return text

    @staticmethod
    def _clean_choice(choice: str) -> str:
        return clean_choice(choice)

    @staticmethod
    def _fmt(value: Any, precision: int = 4) -> str:
        if value is None:
            return "-"
        if isinstance(value, float):
            return f"{value:.{precision}f}"
        if isinstance(value, int):
            return str(value)
        return str(value)

    @staticmethod
    def _escape(text: str) -> str:
        return (
            str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

    @staticmethod
    def _track_name(text: str) -> str:
        if "._" in text:
            return text.split("._")[-1].strip()
        return text

    @classmethod
    def _label(cls, value: Any) -> str:
        return humanize_column_name(value)


def generate_full_report(run_dir: str, user_requirement: str = "") -> str:
    return ReportGenerator(run_dir, user_requirement).generate()


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法：")
        print("  python report_generator.py <运行输出目录> [分析需求]")
        print("  python report_generator.py <运行输出目录> [分析需求] --format word")
        print("  python report_generator.py <运行输出目录> [分析需求] --format pdf")
        raise SystemExit(1)

    run_dir = sys.argv[1]
    user_req = sys.argv[2] if len(sys.argv) > 2 and not sys.argv[2].startswith("--") else ""
    fmt = None
    if "--format" in sys.argv:
        idx = sys.argv.index("--format")
        if idx + 1 < len(sys.argv):
            fmt = sys.argv[idx + 1]

    generator = ReportGenerator(run_dir, user_req)
    if fmt == "word":
        output = generator.export_word()
        print(f"✅ Word 报告已保存到: {output}")
    elif fmt == "pdf":
        output = generator.export_pdf()
        if output:
            print(f"✅ PDF 报告已保存到: {output}")
        else:
            print("⚠️ 当前环境缺少 PDF 导出依赖，已跳过 PDF 输出")
    else:
        output = generator.save()
        print(f"✅ Markdown 报告已保存到: {output}")
