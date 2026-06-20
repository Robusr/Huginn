# -*- coding: utf-8 -*-
"""
@File    : report_generator.py
@Author  : Robusr
@Date    : 2026/6/16
@Description: 完整Markdown报告生成器 — 基于中间JSON结果自动生成7章课程分析报告
@Software: PyCharm
"""

"""
报告生成器
功能：读取 agent_runner.py 生成的标准化中间结果，自动生成域自适应的完整报告。
用法：
    python report_generator.py outputs/<run_dir> "用户分析需求"
    python report_generator.py outputs/<run_dir> "用户分析需求" --format word
"""

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from config import Config, clean_field_name
from logger import get_logger

logger = get_logger(__name__)


class ReportGenerator:
    """域自适应报告生成器：读取中间 JSON → 组装完整 Markdown 报告。"""

    # 默认章节（不含业务模块章）
    _BASE_CHAPTERS = [
        "一、数据来源与分析目标",
        "二、数据概况与预处理",
        "三、描述性统计与可视化",
        "四、统计推断分析",
        "五、业务模块分析",
        "六、主要数据发现",
        "七、改进建议",
        "八、局限性说明",
    ]

    # 图表中文描述映射
    CHART_LABELS = Config.CHART_LABELS

    def __init__(
        self,
        run_dir: Union[str, Path],
        user_requirement: str = "",
        domain_config=None,
    ) -> None:
        """
        初始化报告生成器。
        :param run_dir: agent_runner.py 输出的运行目录路径
        :param user_requirement: 用户输入的分析需求
        :param domain_config: 领域配置（DomainConfig）
        """
        self.run_dir = Path(run_dir)
        self.user_requirement = user_requirement
        self.domain_config = domain_config

        # 域自适应章节
        if domain_config and domain_config.key == "education_survey":
            self.CHAPTERS = [
                "一、数据来源与分析目标",
                "二、数据概况与预处理",
                "三、描述性统计与可视化",
                "四、统计推断分析",
                "五、主要数据发现",
                "六、课程改进建议",
                "七、局限性说明",
            ]
            self.has_business_chapter = False
        else:
            self.CHAPTERS = self._BASE_CHAPTERS
            self.has_business_chapter = True

        # 数据容器
        self.data_profile: Dict[str, Any] = {}
        self.stats_results: Dict[str, Any] = {}
        self.findings: List[Dict] = []
        self.suggestions: List[Dict] = []
        self.valid_tasks: List[Dict] = []
        self.validation_result: Dict[str, Any] = {}
        self.chart_files: List[str] = []
        self.evidence_table: Dict[str, Any] = {}
        self.loss_driver: Dict[str, Any] = {}
        self.discount_analysis: Dict[str, Any] = {}
        self.pareto_results: Dict[str, Any] = {}
        self.cross_dim_results: Dict[str, Any] = {}
        self.granularity: Dict[str, Any] = {}
        self.field_registry: Dict[str, Any] = {}

        # 报告生成时间
        self.generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    @property
    def is_education(self) -> bool:
        return self.domain_config is not None and self.domain_config.key == "education_survey"

    # ==================================================================
    # 公开 API
    # ==================================================================

    def generate(self) -> str:
        """主入口：加载数据 → 构建 7 章 + 附录 → 返回完整 Markdown 字符串。"""
        self._load_all_data()
        self._discover_charts()
        lines = self._render_report()
        return "\n".join(lines)

    def save(self, filename: str = "final_report.md") -> Path:
        """生成报告并保存到 run_dir。"""
        content = self.generate()
        path = self.run_dir / filename
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info("报告已保存至 %s", path)
        return path

    def export_word(self, filename: str = "final_report.docx") -> Path:
        """导出报告为 Word 文档（需 python-docx）。"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")

        doc = Document()

        # 设置默认字体（尝试中文字体）
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimHei"
        font.size = Pt(11)

        # 获取报告内容，按行解析
        markdown_content = self.generate()
        lines = markdown_content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 一级标题 #
            if line.startswith("# ") and not line.startswith("## "):
                heading_text = line[2:].strip()
                doc.add_heading(heading_text, level=1)

            # 二级标题 ##
            elif line.startswith("## "):
                heading_text = line[3:].strip()
                doc.add_heading(heading_text, level=2)

            # 三级标题 ###
            elif line.startswith("### "):
                heading_text = line[4:].strip()
                doc.add_heading(heading_text, level=3)

            # Markdown 图片 ![desc](path)
            elif line.startswith("![") and "](" in line:
                match = re.match(r"!\[(.+?)\]\((.+?)\)", line)
                if match:
                    desc, img_path = match.groups()
                    full_path = self.run_dir / img_path
                    if full_path.exists():
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(desc)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        try:
                            doc.add_picture(str(full_path), width=Inches(5.5))
                        except Exception:
                            p2 = doc.add_paragraph(f"[图片: {desc}]")
                            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 水平线 ---
            elif line == "---":
                doc.add_paragraph("─" * 60)

            # 块引用 >
            elif line.startswith(">"):
                p = doc.add_paragraph()
                run = p.add_run(line[1:].strip())
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)

            # 表格行 |
            elif line.startswith("|") and line.endswith("|"):
                # 跳过分隔行（如 |---|---|）
                if re.match(r"^\|[\s\-:]+\|", line):
                    continue
                # 将表格行作为普通段落处理（完整的 docx 表格转换过于复杂）
                cells = [c.strip() for c in line.split("|")[1:-1]]
                p = doc.add_paragraph("  |  ".join(cells))

            # 列表项 - 或 *
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")

            # 数字列表
            elif re.match(r"^\d+\.\s", line):
                doc.add_paragraph(line, style="List Number")

            # 普通段落
            else:
                # 处理加粗 **text**
                p = doc.add_paragraph()
                parts = re.split(r"(\*\*.+?\*\*)", line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

        path = self.run_dir / filename
        doc.save(str(path))
        logger.info("Word 报告已保存至 %s", path)
        return path

    # ==================================================================
    # 数据加载
    # ==================================================================

    def _load_all_data(self) -> None:
        """加载运行目录中的所有中间结果文件。"""
        # 必需文件
        self.data_profile = self._load_json("data_profile.json", required=True)
        self.stats_results = self._load_json("stats_results.json", required=True)

        # 可选文件
        self.findings = self._load_json_list("findings.json")
        self.suggestions = self._load_json_list("suggestions.json")
        self.valid_tasks = self._load_json_list("valid_tasks.json")
        self.validation_result = self._load_json("validation_result.json", required=False)
        self.evidence_table = self._load_json("evidence_table.json", required=False)
        self.field_registry = self._load_json("field_registry.json", required=False)
        self.granularity = self._load_json("granularity.json", required=False)
        self.loss_driver = self._load_json("loss_driver_results.json", required=False)
        self.discount_analysis = self._load_json("discount_analysis_results.json", required=False)
        self.pareto_results = self._load_json("pareto_results.json", required=False)
        self.cross_dim_results = self._load_json("cross_dimension_results.json", required=False)

    def _load_json(self, filename: str, required: bool = False) -> Dict[str, Any]:
        """加载单个 JSON 文件。"""
        path = self.run_dir / filename
        if not path.exists():
            if required:
                raise FileNotFoundError(f"缺少必需文件: {filename}（路径: {self.run_dir}）")
            return {}
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    def _load_json_list(self, filename: str) -> List[Dict]:
        """加载 JSON 数组文件。"""
        path = self.run_dir / filename
        if not path.exists():
            return []
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                return data
            return []

    def _discover_charts(self) -> None:
        """发现 charts/ 目录下的所有 PNG 图表，按固定顺序排序。"""
        chart_dir = self.run_dir / "charts"
        if not chart_dir.exists():
            return

        png_files = sorted(chart_dir.glob("*.png"))
        # 固定顺序：bar → box → scatter → heatmap → 其他
        preferred = ["bar_chart", "box_plot", "scatter_plot", "correlation_heatmap"]
        ordered = []
        for name in preferred:
            for f in png_files:
                if f.stem == name:
                    ordered.append(f"charts/{f.name}")
                    break
        # 追加未在优先列表中的
        for f in png_files:
            rel = f"charts/{f.name}"
            if rel not in ordered:
                ordered.append(rel)

        self.chart_files = ordered

    # ==================================================================
    # 报告渲染 — 主入口
    # ==================================================================

    def _render_report(self) -> List[str]:
        """渲染完整报告。"""
        lines: List[str] = []

        # 报告头
        lines.extend(self._render_header())

        # 目录
        lines.extend(self._render_toc())

        # 各章节
        lines.extend(self._render_chapter_1())
        lines.extend(self._render_chapter_2())
        lines.extend(self._render_chapter_3())
        lines.extend(self._render_chapter_4())
        if self.has_business_chapter:
            lines.extend(self._render_chapter_5_business())
        lines.extend(self._render_chapter_findings())
        lines.extend(self._render_chapter_suggestions())
        lines.extend(self._render_chapter_limitations())

        # 附录
        lines.extend(self._render_appendix())

        # 页脚
        lines.extend(self._render_footer())

        return lines

    # ==================================================================
    # 报告头
    # ==================================================================

    def _render_header(self) -> List[str]:
        profile_meta = self.data_profile.get("meta", {})
        n_rows = profile_meta.get("n_rows", "?")
        n_cols = profile_meta.get("n_columns", "?")

        # 域自适应标题
        if self.domain_config:
            report_title = self.domain_config.report_title
            report_subtitle = self.domain_config.report_subtitle
        else:
            fp_stem = self.run_dir.name.split("_", 1)[-1] if "_" in self.run_dir.name else ""
            report_title = f"{fp_stem}数据分析报告" if fp_stem else "数据分析报告"
            report_subtitle = "基于数据探索的分析与建议"

        lines = [
            f"# {report_title}",
            "",
            f"> **生成时间**：{self.generated_at}",
            f"> **数据规模**：{n_rows} 行 × {n_cols} 列",
        ]
        if self.user_requirement:
            lines.append(f"> **分析需求**：{self.user_requirement}")
        lines.append(f"> **报告版本**：Huginn {Config.APP_VERSION} 自动生成")
        lines.extend(["", "---", ""])

        # 执行摘要
        lines.extend(self._render_executive_summary())
        return lines

    def _render_executive_summary(self) -> List[str]:
        """生成执行摘要，快速概览报告核心发现。"""
        lines = ["## 执行摘要", ""]

        meta = self.data_profile.get("meta", {})
        n_rows = meta.get("n_rows", "?")

        # 粒度信息
        gran_note = ""
        if self.granularity:
            ent = self.granularity.get("entity_description", "")
            if ent:
                gran_note = f"（{ent}）"
        lines.append(f"本报告基于 **{n_rows}** 条数据记录{gran_note}。")

        # 显著结果统计
        sig = self._extract_significant_results()
        total_sig = len(sig.get("hypothesis", [])) + len(sig.get("anova", [])) + len(sig.get("chi_square", []))

        if total_sig > 0:
            lines.append(f"共发现 **{total_sig}** 个统计显著结果（p<0.05），涵盖组间差异、分类关联和分布特征。")
        else:
            lines.append("本次分析未发现统计显著（p<0.05）的结果，以下基于描述性统计和探索性分析。")

        # 验证得分
        val_score = self.validation_result.get("meta", {}).get("score")
        if val_score is not None:
            val_pass = self.validation_result.get("meta", {}).get("overall_pass", False)
            status = "✅" if val_pass else "⚠️"
            lines.append(f"合规性: {status} {val_score}/100")

        lines.extend(["", "---", ""])
        return lines

    def _render_toc(self) -> List[str]:
        lines = ["## 目录", ""]
        for i, chapter in enumerate(self.CHAPTERS, 1):
            lines.append(f"{i}. [{chapter}](#{chapter.replace('、', '').replace(' ', '-')})")
        lines.append(f"8. [附录：合规性验证报告](#附录合规性验证报告)")
        lines.extend(["", "---", ""])
        return lines

    # ==================================================================
    # 第一章：数据来源与分析目标
    # ==================================================================

    def _render_chapter_1(self) -> List[str]:
        profile_meta = self.data_profile.get("meta", {})
        n_rows = profile_meta.get("n_rows", "未知")
        n_cols = profile_meta.get("n_columns", "未知")

        lines = [
            "# 一、数据来源与分析目标",
            "",
            "## 1.1 数据来源",
            "",
        ]

        if self.is_education:
            lines.append(f"本报告所分析的数据来源于课程问卷调研，共收集有效样本 **{n_rows}** 份，"
                         f"涵盖 **{n_cols}** 个调查维度。")
        elif self.granularity:
            lines.append(f"本报告基于 **{n_rows}** 条数据记录进行分析。")
            lines.append(f"**数据粒度**：{self.granularity.get('entity_description', '明细行')}")
            lines.append(f"**唯一订单数**：{self.granularity.get('unique_order_ids', '未知')}")
            lines.append(f"**唯一客户数**：{self.granularity.get('unique_customer_ids', '未知')}")
        else:
            lines.append(f"本报告基于 **{n_rows}** 条数据记录、**{n_cols}** 个字段进行分析。")

        lines.extend([
            "",
            "## 1.2 分析目标",
            "",
        ])

        if self.user_requirement:
            lines.append(f"核心目标：**{self.user_requirement}**。")
        else:
            lines.append("核心目标：对数据进行探索型统计分析，发现模式、诊断问题、提出改进建议。")

        lines.extend([
            "",
            f"共执行 **{len(self.valid_tasks)}** 项统计分析任务，涵盖以下方法：",
            "",
        ])

        method_counts: Dict[str, int] = {}
        for task in self.valid_tasks:
            m = task.get("method", "其他")
            method_counts[m] = method_counts.get(m, 0) + 1
        for method, count in sorted(method_counts.items(), key=lambda x: -x[1]):
            lines.append(f"- **{method}**：{count} 项")

        lines.extend(["", "---", ""])
        return lines

    # ==================================================================
    # 第二章：数据概况与预处理
    # ==================================================================

    def _render_chapter_2(self) -> List[str]:
        profile_meta = self.data_profile.get("meta", {})
        overview = self.data_profile.get("overview", {})
        fields = self.data_profile.get("fields", [])

        n_rows = profile_meta.get("n_rows", "?")
        n_cols = profile_meta.get("n_columns", "?")
        total_missing_pct = profile_meta.get("total_missing_pct", 0)
        missing = profile_meta.get("total_missing", 0)
        field_types = overview.get("field_type_counts", {})
        dup_rows = overview.get("duplicate_rows", 0)

        lines = [
            "# 二、数据概况与预处理",
            "",
            "## 2.1 数据概览",
            "",
            f"| 指标 | 数值 |",
            f"|------|------|",
            f"| 总行数 | {n_rows} |",
            f"| 总列数 | {n_cols} |",
            f"| 总缺失值 | {missing} |",
            f"| 总体缺失率 | {total_missing_pct:.2f}% |",
            f"| 重复行数 | {dup_rows} |",
            "",
            "## 2.2 字段类型分布",
            "",
            "| 类型 | 数量 |",
            "|------|------|",
        ]
        for ftype, count in sorted(field_types.items(), key=lambda x: -x[1]):
            type_labels = {
                "numeric_continuous": "连续数值型",
                "numeric_discrete": "离散数值型",
                "categorical": "分类型",
                "datetime": "日期时间型",
                "text": "文本型",
                "unknown": "未知类型",
            }
            lines.append(f"| {type_labels.get(ftype, ftype)} | {count} |")

        lines.extend([
            "",
            "## 2.3 字段详细信息",
            "",
            "| 序号 | 字段名 | 推断类型 | 有效值数 | 缺失数 | 缺失率 | 唯一值数 |",
            "|------|--------|----------|----------|--------|--------|----------|",
        ])

        for i, field in enumerate(fields, 1):
            col = field.get("column", "?")
            display_name = clean_field_name(col)
            inferred = field.get("inferred_type", "?")
            count = field.get("count", "?")
            miss = field.get("missing", 0)
            miss_pct = field.get("missing_pct", 0)
            unique = field.get("unique", "?")

            type_labels_short = {
                "numeric_continuous": "连续数值",
                "numeric_discrete": "离散数值",
                "categorical": "分类",
                "datetime": "日期时间",
                "text": "文本",
            }
            type_label = type_labels_short.get(inferred, inferred)

            # 高缺失率标红提示
            warning = " ⚠️" if isinstance(miss_pct, (int, float)) and miss_pct > 30 else ""
            lines.append(
                f"| {i} | {display_name} | {type_label} | {count} | {miss} | "
                f"{miss_pct:.1f}%{warning} | {unique} |"
            )

        lines.extend([
            "",
            "## 2.4 预处理说明",
            "",
            "数据预处理包括以下步骤：",
            "",
            "1. **编码自动识别**：自动探测 CSV 文件编码（UTF-8 → GBK → GB18030），解决中文乱码问题",
            "2. **表头清洗**：去除字段名前后的空白字符，统一格式",
            "3. **空行处理**：删除缺失比例超过 90% 的空行",
            "4. **缺失值处理**：标记缺失值（NA、N/A、null、#N/A 等），在统计计算时自动排除",
            "5. **类型推断**：基于数据特征自动推断字段类型（连续数值 / 离散数值 / 分类 / 文本 / 日期时间）",
            "6. **异常值检测**：使用 IQR 方法标记离群值，但不自动删除，保留在原始数据中",
            "",
            "---",
            "",
        ])

        return lines

    # ==================================================================
    # 第三章：描述性统计与可视化
    # ==================================================================

    def _render_chapter_3(self) -> List[str]:
        fields = self.data_profile.get("fields", [])
        lines = [
            "# 三、描述性统计与可视化",
            "",
            "## 3.1 数值型字段描述统计",
            "",
        ]

        numeric_fields = [f for f in fields if f.get("inferred_type", "").startswith("numeric")]
        categorical_fields = [f for f in fields if f.get("inferred_type") == "categorical"]

        # 元数据字段（不应在分类频数中详细展示）
        _meta_columns = {"提交答卷时间", "所用时间", "来源详情", "序号"}

        if numeric_fields:
            lines.extend([
                "| 字段名 | 样本量 | 均值 | 中位数 | 标准差 | 最小值 | 最大值 | 偏度 | 峰度 |",
                "|--------|--------|------|--------|--------|--------|--------|------|------|",
            ])
            for field in numeric_fields:
                stats = field.get("stats", {})
                if not stats or "error" in field:
                    continue
                col = field.get("column", "?")
                display_name = clean_field_name(col)
                # 兼容新旧 stats 键名
                n = stats.get("n", stats.get("n_total", "?"))
                mean = self._fmt(stats.get("mean"))
                median = self._fmt(stats.get("median"))
                std = self._fmt(stats.get("std"))
                vmin = self._fmt(stats.get("min"))
                vmax = self._fmt(stats.get("max"))
                skew = self._fmt(stats.get("skewness"))
                kurt = self._fmt(stats.get("kurtosis"))
                lines.append(
                    f"| {display_name} | {n} | {mean} | {median} | {std} | "
                    f"{vmin} | {vmax} | {skew} | {kurt} |"
                )
            lines.append("")
        else:
            lines.append("*无连续数值型字段可供描述统计。*")
            lines.append("")

        # 分类字段频数分布（跳过元数据字段，最多展示 8 个）
        if categorical_fields:
            lines.extend([
                "## 3.2 分类型字段频数分布",
                "",
            ])
            shown = 0
            for field in categorical_fields:
                col = field.get("column", "?")
                display_name = clean_field_name(col)
                # 跳过元数据和技术ID字段
                if col in _meta_columns:
                    continue
                if shown >= 8:
                    break
                stats = field.get("stats", {})
                freq = stats.get("frequency_distribution", {})

                lines.append(f"### {display_name}")
                lines.append("")
                lines.append("| 类别 | 频数 | 占比 |")
                lines.append("|------|------|------|")

                for cat, info in list(freq.items())[:10]:  # 最多10个类别
                    lines.append(
                        f"| {cat} | {info.get('count', '?')} | "
                        f"{info.get('pct', 0):.1f}% |"
                    )
                lines.append("")
                shown += 1

        # 图表展示
        lines.extend([
            "## 3.3 可视化图表",
            "",
        ])

        if self.chart_files:
            for chart in self.chart_files:
                chart_stem = Path(chart).stem
                desc = self.CHART_LABELS.get(chart_stem, chart_stem)
                lines.append(f"### {desc}")
                lines.append("")
                lines.append(f"![{desc}]({chart})")
                lines.append("")

            # 图表解读
            lines.extend([
                "### 图表解读说明",
                "",
                "- **柱状图**：展示分类变量的频数分布和分组均值对比，直观反映不同类别的样本量和平均水平差异",
                "- **箱线图**：展示数值变量的分布特征（中位数、四分位数、范围），并标记离群值",
                "- **散点图**：展示两个数值变量之间的关系形态，含回归趋势线和 Pearson 相关系数",
                "- **相关性热力图**：以颜色深浅展示多个数值变量之间的相关强度和方向",
                "",
            ])
        else:
            lines.append("*未生成可视化图表（可能因为数据中无数值列或样本量不足）。*")
            lines.append("")

        lines.extend(["---", ""])
        return lines

    # ==================================================================
    # 第四章：统计推断分析（仅显著结果）
    # ==================================================================

    def _render_chapter_4(self) -> List[str]:
        sig = self._extract_significant_results()
        lines = [
            "# 四、统计推断分析",
            "",
            "> 本章仅展示 **p < 0.05** 的统计显著结果。完整的统计结果请参见 `stats_results.json`。",
            "",
        ]

        # 点估计摘要
        lines.extend(self._render_point_estimation_section())

        # 区间估计摘要
        lines.extend(self._render_interval_estimation_section())

        # 假设检验显著结果
        lines.extend(self._render_significant_hypothesis_tests(sig.get("hypothesis", [])))

        # ANOVA 显著结果
        lines.extend(self._render_significant_anova(sig.get("anova", [])))

        # 卡方检验显著结果
        lines.extend(self._render_significant_chi_square(sig.get("chi_square", [])))

        # 分布检验结果
        lines.extend(self._render_distribution_tests())

        lines.extend(["---", ""])
        return lines

    def _render_point_estimation_section(self) -> List[str]:
        pe = self.stats_results.get("point_estimation", {}).get("fields", {})
        if not pe:
            return ["## 4.1 点估计", "", "*无点估计数据。*", ""]

        lines = [
            "## 4.1 点估计",
            "",
            f"共对 **{len(pe)}** 个数值列执行了点估计，每列计算 10 个参数（均值、方差、标准差、偏度、峰度、变异系数、中位数、IQR、极差、标准误）。",
            "",
            "| 字段 | 样本量 | 均值 | 标准差 | 中位数 | IQR | 偏度 | 峰度 | 变异系数(%) |",
            "|------|--------|------|--------|--------|-----|------|------|-------------|",
        ]
        for col, info in pe.items():
            if not isinstance(info, dict) or "error" in info:
                continue
            display_name = clean_field_name(col)
            n = info.get("n", "?")
            mean = self._fmt(info.get("mean"))
            std = self._fmt(info.get("std"))
            med = self._fmt(info.get("median"))
            iqr = self._fmt(info.get("iqr"))
            skew = self._fmt(info.get("skewness"))
            kurt = self._fmt(info.get("kurtosis_excess"))
            cv = f"{info['cv_pct']:.1f}" if info.get("cv_pct") is not None else "-"
            lines.append(
                f"| {display_name} | {n} | {mean} | {std} | {med} | "
                f"{iqr} | {skew} | {kurt} | {cv} |"
            )
        lines.append("")
        return lines

    def _render_interval_estimation_section(self) -> List[str]:
        ie = self.stats_results.get("interval_estimation", {}).get("fields", {})
        if not ie:
            return ["## 4.2 区间估计", "", "*无区间估计数据。*", ""]

        lines = [
            "## 4.2 区间估计",
            "",
            f"共对 **{len(ie)}** 个数值列执行了 95% 置信水平的区间估计，每列计算 5 类置信区间。",
            "",
            "| 字段 | 样本量 | 均值 95% CI | 标准差 95% CI |",
            "|------|--------|-------------|---------------|",
        ]
        for col, info in ie.items():
            if not isinstance(info, dict) or "error" in info:
                continue
            display_name = clean_field_name(col)
            n = info.get("n", "?")
            mean_ci = info.get("mean_ci")
            std_ci = info.get("std_ci")
            mean_ci_str = f"[{mean_ci[0]:.4f}, {mean_ci[1]:.4f}]" if mean_ci and len(mean_ci) == 2 else "-"
            std_ci_str = f"[{std_ci[0]:.4f}, {std_ci[1]:.4f}]" if std_ci and len(std_ci) == 2 else "-"
            lines.append(f"| {display_name} | {n} | {mean_ci_str} | {std_ci_str} |")
        lines.append("")
        return lines

    def _render_significant_hypothesis_tests(self, sig_tests: List[Dict]) -> List[str]:
        lines = [
            "## 4.3 假设检验（仅显著结果）",
            "",
        ]

        if not sig_tests:
            lines.append("*无统计显著的假设检验结果。*")
            lines.append("")
            return lines

        lines.extend([
            f"共发现 **{len(sig_tests)}** 个统计显著的假设检验结果：",
            "",
            "| # | 方法 | 变量/对比 | 统计量 | p 值 | 结论 |",
            "|---|------|-----------|--------|------|------|",
        ])

        for i, item in enumerate(sig_tests, 1):
            method = item.get("method", "?")
            var_info = item.get("variables", item.get("column", "?"))
            stat = self._fmt(item.get("statistic"))
            p_val = item.get("p_value", "?")
            p_str = f"**{p_val:.4f}**" if isinstance(p_val, (int, float)) else str(p_val)

            # 构造结论
            if isinstance(p_val, (int, float)):
                if p_val < 0.01:
                    conclusion = "差异极显著（p < 0.01）"
                elif p_val < 0.05:
                    conclusion = "差异显著（p < 0.05）"
                else:
                    conclusion = f"不显著（p = {p_val:.4f}）"
            else:
                conclusion = "-"

            lines.append(f"| {i} | {method} | {var_info} | {stat} | {p_str} | {conclusion} |")

        lines.append("")
        return lines

    def _render_significant_anova(self, sig_anova: List[Dict]) -> List[str]:
        lines = [
            "## 4.4 方差分析（仅显著结果）",
            "",
        ]

        if not sig_anova:
            lines.append("*无统计显著的 ANOVA 结果。*")
            lines.append("")
            return lines

        lines.extend([
            f"共发现 **{len(sig_anova)}** 个统计显著的 ANOVA 结果：",
            "",
            "| # | 因变量 | 因子 | 组数 | F 值 | p 值 | 结论 |",
            "|---|--------|------|------|------|------|------|",
        ])

        for i, item in enumerate(sig_anova, 1):
            dep = item.get("dependent", "?")
            factor = item.get("factor", "?")
            n_groups = item.get("n_groups", "?")
            f_stat = self._fmt(item.get("F_statistic"))
            p_val = item.get("p_value", "?")
            p_str = f"**{p_val:.4f}**" if isinstance(p_val, (int, float)) else str(p_val)
            conclusion = f"{factor}对{dep}存在显著主效应（p = {p_str}）"

            lines.append(f"| {i} | {dep} | {factor} | {n_groups} | {f_stat} | {p_str} | {conclusion} |")

            # 如果有 Tukey HSD 事后检验结果，简要说明
            tukey = item.get("tukey_hsd")
            if tukey:
                lines.append(f"| | | > 事后检验（Tukey HSD）已执行，详情见 `stats_results.json` |")
        lines.append("")
        return lines

    def _render_significant_chi_square(self, sig_chi: List[Dict]) -> List[str]:
        lines = [
            "## 4.5 卡方检验（仅显著结果）",
            "",
        ]

        if not sig_chi:
            lines.append("*无统计显著的卡方检验结果。*")
            lines.append("")
            return lines

        lines.extend([
            f"共发现 **{len(sig_chi)}** 个统计显著的卡方检验结果：",
            "",
            "| # | 检验变量 | χ² 值 | 自由度 | p 值 | 结论 |",
            "|---|----------|-------|--------|------|------|",
        ])

        for i, item in enumerate(sig_chi, 1):
            var_info = item.get("variables", item.get("column", "?"))
            chi2 = self._fmt(item.get("chi2_statistic"))
            df = item.get("df", "?")
            p_val = item.get("p_value", "?")
            p_str = f"**{p_val:.4f}**" if isinstance(p_val, (int, float)) else str(p_val)

            if isinstance(p_val, (int, float)) and p_val < 0.05:
                conclusion = "各类别分布存在显著差异"
            else:
                conclusion = f"不显著"

            lines.append(f"| {i} | {var_info} | {chi2} | {df} | {p_str} | {conclusion} |")
        lines.append("")
        return lines

    def _render_distribution_tests(self) -> List[str]:
        dist = self.stats_results.get("distribution_tests", {}).get("tests", {})
        if not dist:
            return ["## 4.6 分布检验", "", "*无分布检验数据。*", ""]

        lines = [
            "## 4.6 分布检验（正态性检验）",
            "",
            "| 字段 | 样本量 | Shapiro-Wilk W | Shapiro-Wilk p | D'Agostino χ² | D'Agostino p | 正态性 |",
            "|------|--------|----------------|----------------|----------------|---------------|--------|",
        ]

        for col, info in dist.items():
            if not isinstance(info, dict) or "error" in info:
                continue
            display_name = clean_field_name(col)
            n = info.get("n", "?")

            sw = info.get("shapiro_wilk", {})
            sw_stat = self._fmt(sw.get("statistic")) if isinstance(sw, dict) else "-"
            sw_p = sw.get("p_value") if isinstance(sw, dict) else None
            sw_p_str = f"{sw_p:.4f}" if isinstance(sw_p, (int, float)) else "-"

            dp = info.get("dagostino_pearson", {})
            dp_stat = self._fmt(dp.get("statistic")) if isinstance(dp, dict) else "-"
            dp_p = dp.get("p_value") if isinstance(dp, dict) else None
            dp_p_str = f"{dp_p:.4f}" if isinstance(dp_p, (int, float)) else "-"

            # 正态性判断
            if isinstance(sw_p, (int, float)) and isinstance(dp_p, (int, float)):
                if sw_p >= 0.05 and dp_p >= 0.05:
                    normality = "✅ 正态"
                elif sw_p < 0.05 or dp_p < 0.05:
                    normality = "⚠️ 非正态"
                else:
                    normality = "-"
            else:
                normality = "-"

            lines.append(
                f"| {display_name} | {n} | {sw_stat} | {sw_p_str} | "
                f"{dp_stat} | {dp_p_str} | {normality} |"
            )

        lines.extend([
            "",
            "> **说明**：Shapiro-Wilk 检验适用于样本量 3 ≤ n ≤ 5000，" "",
            "> D'Agostino-Pearson 检验适用于较大样本量。",
            "> 标注 ⚠️ 的字段分布显著偏离正态，分析时应优先使用非参数方法（如 Wilcoxon、Mann-Whitney U）。",
            "",
        ])
        return lines

    # ==================================================================
    # 第五章（业务模块分析）
    # ==================================================================

    def _render_chapter_5_business(self) -> List[str]:
        """渲染业务模块分析章节（零售等非教育域）。"""
        lines = [
            "# 五、业务模块分析",
            "",
        ]

        has_content = False

        # 5.1 数据粒度
        if self.granularity:
            has_content = True
            lines.extend([
                "## 5.1 数据粒度与实体识别",
                "",
                f"- **行级实体**：{self.granularity.get('entity_description', '未知')}",
                f"- **总记录数**：{self.granularity.get('row_count', '?')}",
                f"- **唯一订单数**：{self.granularity.get('unique_order_ids', '?')}",
                f"- **唯一客户数**：{self.granularity.get('unique_customer_ids', '?')}",
                f"- **唯一产品数**：{self.granularity.get('unique_product_ids', '?')}",
                "",
                "> ⚠️ **重要**：以下所有比率分析均已明确分母。亏损明细率 ≠ 亏损订单率 ≠ 亏损客户率。",
                "",
            ])

        # 5.2 亏损驱动分析
        loss = self.loss_driver
        if loss and loss.get("is_viable"):
            has_content = True
            lines.extend([
                "## 5.2 亏损驱动分析",
                "",
            ])

            # 总体指标
            ov = loss.get("overall_summary", {})
            if ov:
                lines.extend([
                    "| 指标 | 数值 |",
                    "|------|------|",
                    f"| 总销售额 | {ov.get('total_sales', '?'):,.2f} |",
                    f"| 总利润 | {ov.get('total_profit', '?'):,.2f} |",
                    f"| 总体亏损率（明细行） | {ov.get('overall_loss_rate', 0)*100:.2f}% |",
                    f"| 总体利润率 | {ov.get('overall_profit_margin', '?'):.2f}% |",
                    "",
                ])

            # 主要亏损来源（合并所有维度）
            top_contributors = loss.get("top_loss_contributors", [])
            if top_contributors:
                lines.extend([
                    "### 主要亏损来源（按亏损贡献率排序）",
                    "",
                    "| 维度 | 分类 | 销售额 | 利润 | 利润率 | 亏损率 | 亏损金额 | 亏损贡献% | 平均折扣 |",
                    "|------|------|--------|------|--------|--------|----------|-----------|----------|",
                ])
                for item in top_contributors[:10]:
                    dim = item.get("dimension_display", "")
                    name = item.get("name", "")
                    sales = item.get("total_sales", 0)
                    profit = item.get("total_profit", 0)
                    margin = item.get("profit_margin", 0)
                    loss_rate = item.get("loss_rate", 0) * 100
                    loss_amt = item.get("loss_amount", 0)
                    loss_pct = item.get("loss_contribution_pct", 0)
                    avg_disc = item.get("avg_discount", 0)
                    disc_str = f"{avg_disc*100:.1f}%" if avg_disc else "-"
                    lines.append(
                        f"| {dim} | {name} | {sales:,.0f} | {profit:,.0f} | "
                        f"{margin:.1f}% | {loss_rate:.1f}% | {loss_amt:,.0f} | "
                        f"{loss_pct:.1f}% | {disc_str} |"
                    )
                lines.append("")

                # 亏损集中度摘要
                dim_results = loss.get("dimension_results", {})
                for dim_name, dim_result in dim_results.items():
                    if isinstance(dim_result, dict) and "summary" in dim_result:
                        summary = dim_result["summary"]
                        name = summary.get("largest_loss_contributor", {})
                        if name:
                            lines.append(
                                f"- **{dim_result.get('dimension_display', dim_name)}**："
                                f"最大亏损来源为「{name.get('name', '?')}」"
                                f"（亏损贡献 {name.get('loss_contribution_pct', 0):.1f}%），"
                                f"{summary.get('loss_concentration', '')}"
                            )
                lines.append("")

        # 5.3 折扣响应分析
        disc = self.discount_analysis
        if disc and disc.get("is_viable"):
            has_content = True
            lines.extend([
                "## 5.3 折扣响应分析",
                "",
            ])

            # 总体折扣统计
            ov = disc.get("overall_summary", {})
            if ov:
                lines.extend([
                    f"- 平均折扣率：{ov.get('mean_discount', 0)*100:.1f}%",
                    f"- 中位数折扣率：{ov.get('median_discount', 0)*100:.1f}%",
                    f"- 打折交易占比：{ov.get('discount_rate', 0)*100:.1f}%",
                    "",
                ])

            # 折扣分箱
            bins = disc.get("discount_bins", {}).get("bins", [])
            if bins:
                lines.extend([
                    "### 折扣分箱分析",
                    "",
                    "| 折扣区间 | 交易数 | 总销售额 | 平均销售额 | 总利润 | 平均利润 |",
                    "|----------|--------|----------|------------|--------|----------|",
                ])
                for b in bins:
                    lines.append(
                        f"| {b.get('bin', '?')} | {b.get('count', 0)} | "
                        f"{b.get('total_sales', 0):,.0f} | {b.get('avg_sales', 0):,.0f} | "
                        f"{b.get('total_profit', 0):,.0f} | {b.get('avg_profit', 0):,.0f} |"
                    )
                lines.append("")

            # 利润率阈值
            tp = disc.get("profit_tipping_point", {})
            if tp and tp.get("tipping_bin"):
                lines.extend([
                    "### 折扣阈值分析",
                    "",
                    f"**{tp.get('description', '')}**",
                    "",
                ])

            # 异常检测
            anomalies = disc.get("anomalies", {})
            if anomalies and anomalies.get("anomaly_count", 0) > 0:
                lines.extend([
                    "### 高折扣异常检测",
                    "",
                    f"{anomalies.get('summary', '')}",
                    "",
                ])

        # 5.4 集中度与帕累托分析
        pareto = self.pareto_results
        if pareto and pareto.get("is_viable"):
            has_content = True
            lines.extend([
                "## 5.4 集中度与帕累托分析",
                "",
            ])

            for key in ["product_concentration", "customer_concentration", "subcategory_concentration"]:
                conc = pareto.get(key, {})
                if not conc or "error" in conc:
                    continue
                label = conc.get("label", key)
                metrics = conc.get("concentration_metrics", {})
                lines.extend([
                    f"### {label}集中度",
                    "",
                    f"- 前5个{label}贡献: **{metrics.get('top5_pct', 0):.1f}%** 的销售额",
                    f"- 前20%的{label}贡献: **{metrics.get('top20_pct', 0):.1f}%**",
                    f"- 覆盖80%销售额需要: **{metrics.get('items_needed_for_80pct', '?')}** 个{label}",
                    f"  （占全部{conc.get('total_items', '?')}个{label}的{metrics.get('pct_items_for_80pct', 0):.1f}%）",
                    "",
                ])
                # Top-N表
                top_items = conc.get("top_items", [])[:10]
                if top_items:
                    lines.extend([
                        "| 排名 | 名称 | 销售额 | 占总比% | 累计% | 利润 | 利润率% |",
                        "|------|------|--------|---------|-------|------|---------|",
                    ])
                    for item in top_items:
                        lines.append(
                            f"| {item.get('rank', '?')} | {item.get('name', '?')[:30]} | "
                            f"{item.get('total_value', 0):,.0f} | {item.get('pct_of_total', 0):.1f}% | "
                            f"{item.get('cumulative_pct', 0):.1f}% | "
                            f"{item.get('total_profit', 0):,.0f} | {item.get('profit_margin', 0):.1f}% |"
                        )
                    lines.append("")

            # 异常对象
            hs_lp = pareto.get("high_sales_low_profit", {})
            if hs_lp and hs_lp.get("count", 0) > 0:
                lines.extend([
                    "### 高销售低利润商品",
                    "",
                    f"{hs_lp.get('summary', '')}",
                    "",
                ])

            ls_hl = pareto.get("low_sales_high_loss", {})
            if ls_hl and ls_hl.get("count", 0) > 0:
                lines.extend([
                    "### 低销售高亏损商品",
                    "",
                    f"{ls_hl.get('summary', '')}",
                    "",
                ])

        # 5.5 交叉维度分析
        cross = self.cross_dim_results
        if cross and cross.get("is_viable"):
            combos = cross.get("combinations", [])
            if combos:
                has_content = True
                lines.extend([
                    "## 5.5 交叉维度洞察",
                    "",
                ])
                for combo in combos:
                    if "error" in combo:
                        continue
                    label = combo.get("label", "")
                    n_patterns = combo.get("n_notable_patterns", 0)
                    lines.append(f"### {label}（{n_patterns} 个显著交互模式）")
                    patterns = combo.get("top_patterns", [])[:8]
                    if patterns:
                        lines.extend([
                            "| 维度1 | 维度2 | 销售额 | 交易数 | 利润率 |",
                            "|-------|-------|--------|--------|--------|",
                        ])
                        for p in patterns:
                            lines.append(
                                f"| {p.get('dim1', '?')} | {p.get('dim2', '?')} | "
                                f"{p.get('total_sales', 0):,.0f} | {p.get('transaction_count', 0)} | "
                                f"{p.get('profit_margin', 0):.1f}% |"
                            )
                    lines.append("")

        if not has_content:
            lines.append("*当前数据不支持业务模块分析，或业务模块分析未生成结果。*")
            lines.append("")

        lines.extend(["---", ""])
        return lines

    # ==================================================================
    # 第五章/第六章：主要数据发现（兼容旧名称）
    # ==================================================================

    def _render_chapter_findings(self) -> List[str]:
        """数据发现章节（域自适应标题）。"""
        return self._render_chapter_5()

    def _render_chapter_suggestions(self) -> List[str]:
        """改进建议章节（域自适应标题）。"""
        return self._render_chapter_6()

    def _render_chapter_limitations(self) -> List[str]:
        """局限性说明章节。"""
        return self._render_chapter_7()

    # 保留旧方法名以保证向后兼容
    def _render_chapter_5(self) -> List[str]:
        chapter_title = "五、课程改进建议" if self.is_education else "六、主要数据发现" if self.has_business_chapter else "五、主要数据发现"
        lines = [
            f"# {chapter_title}",
            "",
        ]

        if not self.findings:
            lines.extend([
                "*未生成 LLM 数据发现（可能为离线模式或 findings.json 缺失）。*",
                "",
                "以下为基于统计结果自动提炼的显著发现：",
                "",
            ])
            # 从统计结果中自动提取
            sig = self._extract_significant_results()
            auto_findings = []
            for item in sig.get("hypothesis", [])[:3]:
                auto_findings.append({
                    "conclusion": f"{item.get('method', '')}表明{item.get('variables', '')}存在显著差异",
                    "evidence": f"p={item.get('p_value', '?'):.4f}" if isinstance(item.get('p_value'), (int, float)) else "",
                    "method": item.get("method", ""),
                    "importance": 3,
                })
            for item in sig.get("anova", [])[:2]:
                auto_findings.append({
                    "conclusion": f"{item.get('factor', '')}对{item.get('dependent', '')}存在显著主效应",
                    "evidence": f"F={item.get('F_statistic', '?')}, p={item.get('p_value', '?'):.4f}" if isinstance(item.get('p_value'), (int, float)) else "",
                    "method": item.get("method", "方差分析"),
                    "importance": 4,
                })
            self.findings = auto_findings

        # 按重要性排序
        sorted_findings = sorted(
            self.findings,
            key=lambda x: x.get("importance", 0) if isinstance(x.get("importance"), (int, float)) else 0,
            reverse=True,
        )

        lines.append(f"共识别 **{len(sorted_findings)}** 条核心数据发现，按重要性从高到低排列：")
        lines.append("")

        for i, finding in enumerate(sorted_findings, 1):
            conclusion = finding.get("conclusion", "?")
            evidence = finding.get("evidence", "")
            method = finding.get("method", "")
            importance = finding.get("importance", "?")

            stars = "⭐" * min(importance, 5) if isinstance(importance, int) else ""
            lines.extend([
                f"### 发现 {i}：{conclusion} {stars}",
                "",
                f"- **重要性**：{importance}/5",
                f"- **方法**：{method}",
                f"- **证据**：{evidence}",
                "",
            ])

        lines.extend(["---", ""])
        return lines

    # ==================================================================
    # 第六章：课程改进建议
    # ==================================================================

    def _render_chapter_6(self) -> List[str]:
        chapter_title = "六、课程改进建议" if self.is_education else "七、改进建议" if self.has_business_chapter else "六、改进建议"
        lines = [
            f"# {chapter_title}",
            "",
        ]

        if not self.suggestions:
            lines.extend([
                "*未生成 LLM 课程建议（可能为离线模式或 suggestions.json 缺失）。*",
                "",
                "建议基于上述数据发现，重点关注以下方向：",
                "",
                "1. 针对差异显著的群体，设计差异化教学策略",
                "2. 针对评分较低的模块，优化教学内容和方法",
                "3. 针对满意度较低的维度，深入调研原因并制定改进方案",
                "",
            ])
            lines.extend(["---", ""])
            return lines

        lines.append(f"基于数据发现，共提出 **{len(self.suggestions)}** 条课程改进建议：")
        lines.append("")

        for i, suggestion in enumerate(self.suggestions, 1):
            sug_text = suggestion.get("suggestion", "")
            evidence = suggestion.get("evidence", "")
            direction = suggestion.get("direction", "")

            lines.extend([
                f"### 建议 {i}：{sug_text}",
                "",
                f"- **数据依据**：{evidence}",
                f"- **改进方向**：{direction}",
                "",
            ])

        lines.extend([
            "---",
            "",
        ])
        return lines

    # ==================================================================
    # 第七章：局限性说明
    # ==================================================================

    def _render_chapter_7(self) -> List[str]:
        profile_meta = self.data_profile.get("meta", {})
        n_rows = profile_meta.get("n_rows", "未知")
        total_missing_pct = profile_meta.get("total_missing_pct", 0)

        chapter_title = "七、局限性说明" if self.is_education else "八、局限性说明" if self.has_business_chapter else "七、局限性说明"
        lines = [
            f"# {chapter_title}",
            "",
            "本报告的分析结果受到以下因素的制约，在解读和使用时需要审慎考虑：",
            "",
            "### 7.1 相关性与因果关系",
            "",
            "> ⚠️ **重要提醒**：本报告中所有统计显著的相关关系**不代表因果关系**。",
            "> 两变量之间存在显著相关，仅表明它们存在统计学上的关联，",
            "> 不能推断其中一个变量是另一个变量的原因。干预建议应结合教学理论和实际情况综合判断。",
            "",
            "### 7.2 样本量限制",
            "",
            f"本数据集共包含 **{n_rows}** 个有效观测。样本量的大小直接影响统计检验的效力：",
            "- 样本量过小可能导致检验效力不足（II 类错误），真实存在的差异未能检测出来",
            "- 样本量过大可能使微小的实际差异也变得统计显著，但实际意义有限",
            "",
            "### 7.3 问卷偏差",
            "",
            "课程问卷数据存在以下固有偏差：",
            "- **自陈报告偏差**：学生的主观评分可能受到社会期望、情绪状态、填写环境等因素影响",
            "- **选择偏差**：填答问卷的学生群体可能不完全代表全体学生",
            "- **量表限制**：Likert 量表数据的等距性假设可能不完全成立，但本报告基于常规做法按等距数据处理",
            "",
            "### 7.4 统计显著与实质显著",
            "",
            "统计显著性（p < 0.05）仅表明观察到的差异不太可能由偶然因素造成，",
            "但不等于该差异在实际教学中具有重要价值。在解读结果时，",
            "应结合效应量（如 Cohen's d、η²）和教学实际综合判断。",
            "",
            f"### 7.5 缺失数据影响",
            "",
            f"本数据集的总体缺失率为 **{total_missing_pct:.2f}%**。",
            "缺失数据可能影响分析结果的代表性，特别是当缺失为非随机缺失（MNAR）时。",
            "本报告采用成对删除（pairwise deletion）处理缺失值。",
            "",
            "### 7.6 泛化性限制",
            "",
            "本报告的结论基于特定课程、特定学期的问卷数据，",
            "其结论的泛化性受限于样本的代表性和数据采集的具体情境。",
            "将结论推广到其他课程或学期时需谨慎。",
            "",
            "---",
            "",
        ]
        return lines

    # ==================================================================
    # 附录：合规性验证报告
    # ==================================================================

    def _render_appendix(self) -> List[str]:
        lines = [
            "# 附录：合规性验证报告",
            "",
        ]

        if not self.validation_result:
            lines.extend([
                "*合规性验证尚未运行。请执行以下命令生成验证结果：*",
                "",
                f"```bash",
                f"python report_validator.py {self.run_dir}",
                f"```",
                "",
                "运行后重新生成报告即可在附录中看到验证得分。",
                "",
                "---",
                "",
            ])
            return lines

        meta = self.validation_result.get("meta", {})
        score = meta.get("score", 0)
        passed = meta.get("overall_pass", False)
        generated = meta.get("generated_at", "?")

        lines.extend([
            f"> **验证时间**：{generated}",
            f"> **总分**：**{score}/100**",
            f"> **整体结果**：{'✅ 通过（≥60分）' if passed else '❌ 不通过（<60分）'}",
            "",
            "## 各模块得分",
            "",
            "| 检查模块 | 得分 | 满分 | 结果 |",
            "|----------|------|------|------|",
        ])

        checks = self.validation_result.get("checks", {})
        module_names = {
            "statistical_quantity": "统计数量硬指标",
            "statistical_validity": "统计结果有效性",
            "findings_compliance": "数据发现合规性",
            "suggestions_reasonableness": "课程建议合理性",
            "report_completeness": "报告完整性",
        }
        max_scores = {
            "statistical_quantity": 40,
            "statistical_validity": 20,
            "findings_compliance": 20,
            "suggestions_reasonableness": 10,
            "report_completeness": 10,
        }

        for key, check in checks.items():
            name = module_names.get(key, key)
            check_score = check.get("score", 0)
            max_score = max_scores.get(key, "?")
            status = "✅ 通过" if check.get("pass") else "❌ 未通过"
            lines.append(f"| {name} | {check_score} | {max_score} | {status} |")

        # 改进建议
        improvements = self.validation_result.get("improvement_suggestions", [])
        if improvements:
            lines.extend([
                "",
                "## 改进建议",
                "",
            ])
            for i, sug in enumerate(improvements, 1):
                lines.append(f"{i}. {sug}")

        lines.extend(["", "---", ""])
        return lines

    # ==================================================================
    # 页脚
    # ==================================================================

    def _render_footer(self) -> List[str]:
        return [
            "",
            "---",
            "",
            f"*本报告由 Huginn 数据分析智能体自动生成（版本 {Config.APP_VERSION}）*",
            f"*生成时间：{self.generated_at}*",
            f"*所有统计量由 Python（scipy + statsmodels）真实计算，可溯源至 stats_results.json*",
            "",
        ]

    # ==================================================================
    # 显著结果提取（复用 insight_generator.py 扫描模式）
    # ==================================================================

    def _extract_significant_results(self) -> Dict[str, List[Dict]]:
        """扫描 stats_results.json，提取所有 p < 0.05 的显著结果。"""
        alpha = 0.05
        sig: Dict[str, List[Dict]] = {
            "hypothesis": [],
            "anova": [],
            "chi_square": [],
            "distribution": [],
        }

        # 1. 扫描假设检验
        ht = self.stats_results.get("hypothesis_tests", {}).get("tests", {})
        for test_group_name, test_group in ht.items():
            if not isinstance(test_group, dict):
                continue
            if "error" in test_group:
                continue

            # 单条目检验（如 welch_ttest, grouped_ttest 等）
            if "p_value" in test_group:
                p_val = test_group["p_value"]
                if isinstance(p_val, (int, float)) and p_val < alpha:
                    item = dict(test_group)
                    item["_source"] = test_group_name
                    sig["hypothesis"].append(item)

            # 多条目检验（如 one_sample_ttest, wilcoxon 等）
            for key, val in test_group.items():
                if isinstance(val, dict) and "p_value" in val:
                    p_val = val["p_value"]
                    if isinstance(p_val, (int, float)) and p_val < alpha:
                        val_copy = dict(val)
                        val_copy["_test_group"] = test_group_name
                        val_copy["_column"] = key
                        sig["hypothesis"].append(val_copy)

        # 2. 扫描 ANOVA
        anova_data = self.stats_results.get("anova", {}).get("tests", {})
        for anova_name, anova_item in anova_data.items():
            if not isinstance(anova_item, dict):
                continue
            if "p_value" in anova_item:
                p_val = anova_item["p_value"]
                if isinstance(p_val, (int, float)) and p_val < alpha:
                    item = dict(anova_item)
                    item["_source"] = anova_name
                    sig["anova"].append(item)

        # 3. 扫描卡方检验
        chi_data = self.stats_results.get("chi_square_goodness_of_fit", {}).get("tests", {})
        for chi_name, chi_item in chi_data.items():
            if isinstance(chi_item, dict) and "p_value" in chi_item:
                p_val = chi_item["p_value"]
                if isinstance(p_val, (int, float)) and p_val < alpha:
                    item = dict(chi_item)
                    item["_source"] = chi_name
                    sig["chi_square"].append(item)

        # 4. 扫描分布检验（非正态 = 显著）
        dist_data = self.stats_results.get("distribution_tests", {}).get("tests", {})
        for col, col_data in dist_data.items():
            if not isinstance(col_data, dict):
                continue
            for test_key in ["shapiro_wilk", "dagostino_pearson"]:
                test_info = col_data.get(test_key)
                if isinstance(test_info, dict) and "p_value" in test_info:
                    p_val = test_info["p_value"]
                    if isinstance(p_val, (int, float)) and p_val < alpha:
                        item = dict(test_info)
                        item["_column"] = col
                        item["_test_type"] = test_key
                        sig["distribution"].append(item)

        # 按 p 值排序
        for key in sig:
            sig[key].sort(
                key=lambda x: x.get("p_value", 1.0) if isinstance(x.get("p_value"), (int, float)) else 1.0
            )

        return sig

    # ==================================================================
    # 工具方法
    # ==================================================================

    @staticmethod
    def _fmt(val: Any, precision: int = 4) -> str:
        """格式化数值，保留指定小数位。"""
        if val is None:
            return "-"
        if isinstance(val, float):
            return f"{val:.{precision}f}"
        return str(val)

# ==================================================================
# 便捷函数
# ==================================================================

def generate_full_report(run_dir: str, user_requirement: str = "") -> str:
    """
    一行调用：生成完整 Markdown 报告。
    :param run_dir: 运行输出目录路径
    :param user_requirement: 用户输入的分析需求
    :return: 完整的 Markdown 报告字符串
    """
    gen = ReportGenerator(run_dir, user_requirement)
    return gen.generate()


# ==================================================================
# 命令行入口
# ==================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法：")
        print("  python report_generator.py <运行输出目录> [分析需求]")
        print("  python report_generator.py <运行输出目录> [分析需求] --format word")
        print("示例：")
        print("  python report_generator.py outputs/20260610_143022_课程问卷 '为下一次上课的老师生成课程建议报告'")
        sys.exit(1)

    run_dir = sys.argv[1]
    user_req = sys.argv[2] if len(sys.argv) > 2 and not sys.argv[2].startswith("--") else ""
    export_word = "--format" in sys.argv and "word" in sys.argv

    try:
        gen = ReportGenerator(run_dir, user_req)
        if export_word:
            path = gen.export_word()
            print(f"✅ Word 报告已保存到: {path}")
        else:
            # 默认保存 Markdown
            path = gen.save()
            print(f"✅ Markdown 报告已保存到: {path}")

            # 也打印摘要
            print(f"\n报告包含 {len(gen.CHAPTERS)} 章 + 附录")
            print(f"数据来源: {len(gen.data_profile.get('fields', []))} 个字段")
            print(f"统计结果: {'已加载' if gen.stats_results else '未加载'}")
            print(f"图表数量: {len(gen.chart_files)} 张")
            print(f"数据发现: {len(gen.findings)} 条")
            print(f"课程建议: {len(gen.suggestions)} 条")
    except Exception as e:
        print(f"\n❌ 报告生成失败: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
